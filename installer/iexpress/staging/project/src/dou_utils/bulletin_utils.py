"""
bulletin_utils.py
Geração de boletins em DOCX, Markdown e HTML com agrupamento por (órgão, tipo_ato)
e sumarização (simples ou avançada) opcional.

Função principal:
  generate_bulletin(result_dict, out_path, kind="docx", summarize=False,
					summarizer=None, keywords=None, max_lines=5, mode="center")
"""

from __future__ import annotations
from abc import ABC, abstractmethod
from typing import Dict, Any, List, Callable, Optional, Tuple
from collections import defaultdict
import html as html_lib
from pathlib import Path

from .log_utils import get_logger

logger = get_logger(__name__)

def _remove_dou_metadata(text: str) -> str:
	"""Remove linhas e trechos típicos de metadados do DOU para não poluir resumos.

	Filtra cabeçalhos como "Diário Oficial da União", "Publicado em:",
	"Edição", "Seção", "Página", "Órgão", "Imprensa Nacional", etc.
	Mantém demais linhas intactas.
	"""
	import re
	if not text:
		return ""

	# Remover tags HTML simples que possam estar presentes
	t = re.sub(r"<[^>]+>", " ", text)
	lines = re.split(r"[\r\n]+", t)
	cleaned: list[str] = []
	for ln in lines:
		low = ln.strip().lower()
		if not low:
			continue
		# Padrões de metadados do DOU a remover
		if re.search(r"\b(di[áa]rio oficial da uni[aã]o|imprensa nacional)\b", low):
			continue
		if re.search(r"\b(publicado em|edi[cç][aã]o|se[cç][aã]o|p[aá]gina|[oó]rg[aã]o)\b", low):
			continue
		if re.search(r"\b(bras[aã]o)\b", low):
			continue
		# Disclaimers e elementos de layout
		if re.search(r"este conte[úu]do n[aã]o substitui", low):
			continue
		if re.search(r"borda do rodap[eé]|logo da imprensa|rodap[eé]", low):
			continue
		cleaned.append(ln)

	return "\n".join(cleaned)



def _split_doc_header(text: str) -> Tuple[Optional[str], str]:
	"""Localiza o cabeçalho do ato em qualquer ponto das primeiras linhas e retorna (header, body).

	- Procura por tokens de tipo de ato (DESPACHO, PORTARIA, RESOLUÇÃO, DECRETO, MENSAGEM, etc.)
	- Quando encontra, retorna da posição do tipo até o fim da linha ou primeiro ponto final.
	- O restante do texto (após o cabeçalho) é retornado como body para uso no resumo.
	"""
	import re
	if not text:
		return None, ""

	# Remover tags HTML simples e normalizar quebras
	raw = re.sub(r"<[^>]+>", " ", text)
	# Limitar a janela de busca para desempenho, mas suficientemente ampla
	raw_window = raw[:4000]
	lines = re.split(r"[\r\n]+", raw_window)
	head_candidates = []
	for ln in lines[:50]:
		s = ln.strip()
		if not s:
			continue
		# Não descarte linhas com metadados — o cabeçalho pode estar nelas
		head_candidates.append(s)
	blob = "\n".join(head_candidates)

	# Alternativas com e sem acentuação
	doc_alt = [
		"PORTARIA CONJUNTA", "INSTRUÇÃO NORMATIVA", "INSTRUCAO NORMATIVA", "DECRETO-LEI",
		"MEDIDA PROVISÓRIA", "MEDIDA PROVISORIA", "ORDEM DE SERVIÇO", "ORDEM DE SERVICO",
		"RESOLUÇÃO", "RESOLUCAO", "DELIBERAÇÃO", "DELIBERACAO", "RETIFICAÇÃO", "RETIFICACAO",
		"COMUNICADO", "MENSAGEM", "EXTRATO", "PARECER", "DESPACHO", "EDITAL", "DECRETO",
		"PORTARIA", "LEI", "ATO", "AVISO"
	]
	doc_alt_sorted = sorted(doc_alt, key=len, reverse=True)

	# Encontrar primeira ocorrência de qualquer tipo dentro do blob
	found = None
	start_idx = None
	upper_blob = blob.upper()
	for dt in doc_alt_sorted:
		i = upper_blob.find(dt)
		if i != -1 and (start_idx is None or i < start_idx):
			start_idx = i
			found = dt
	if start_idx is None:
		# fallback: primeira linha com alta proporção de maiúsculas
		for s in head_candidates:
			letters = [ch for ch in s if ch.isalpha()]
			if not letters:
				continue
			upp = sum(1 for ch in letters if ch.isupper())
			if upp / max(1, len(letters)) >= 0.6 and len(s) >= 10:
				header = s
				# body é o restante após esta linha aproximada
				idx = raw.find(s)
				if idx != -1:
					body = raw[idx + len(s):].lstrip(" \t\r\n-—:")
					return header, body
				return header, raw
		return None, text

	# Determinar fim do cabeçalho: até primeiro ponto final ou quebra de linha subsequente
	after = blob[start_idx:]
	# Tentar englobar múltiplas linhas até primeiro ponto final
	header_lines = []
	remain = after
	consumed = 0
	while True:
		ln = remain.split("\n", 1)[0]
		header_lines.append(ln.strip())
		consumed += len(ln) + 1 if "\n" in remain else len(ln)
		# parar se achou ponto final
		if "." in ln:
			break
		# adicionar próxima linha somente se parecer parte do cabeçalho
		tail = remain[len(ln) + (1 if "\n" in remain else 0):]
		if not tail:
			break
		nxt = tail.split("\n", 1)[0].strip()
		if re.search(r"\b(MENSAGEM\s+N[ºO]|N[ºO]\s+\d|de\s+\d{1,2}\s+de)\b", nxt, flags=re.I):
			remain = tail
			continue
		# se a próxima linha ainda for majoritariamente maiúscula, considerar também
		letters = [ch for ch in nxt if ch.isalpha()]
		upp = sum(1 for ch in letters if ch.isupper()) if letters else 0
		if letters and upp / len(letters) >= 0.6:
			remain = tail
			continue
		break
	header = re.sub(r"\s+", " ", " ".join(header_lines)).strip()

	# Construir body removendo o header da primeira ocorrência no raw original
	# Usar uma busca case-insensitive para achar a mesma fatia
	pattern = re.escape(header[:80])  # usar prefixo para casar de forma robusta
	m_body = re.search(pattern, raw, flags=re.I)
	if m_body:
		body = raw[m_body.end():].lstrip(" \t\r\n-—:")
	else:
		body = raw

	return header, body


def _extract_doc_header_line(it: Dict[str, Any]) -> Optional[str]:
	"""Compat: extrai apenas o header do texto do item, com fallback em campos do item."""
	import re
	text = it.get("texto") or it.get("ementa") or ""
	header, _ = _split_doc_header(text)
	if header:
		return header

	# Fallback para campos do item
	candidates: list[str] = []
	for key in ("titulo", "titulo_listagem", "title_friendly"):
		v = it.get(key)
		if v:
			candidates.append(str(v).strip())
	blob = " ".join(candidates)

	tipo = (it.get("tipo_ato") or "").strip()
	if tipo:
		m = re.search(r"\bN[ºO]\s*[\w\-./]+", blob, flags=re.I)
		if m:
			return f"{tipo.upper()} {m.group(0)}"[:200]
		return tipo.upper()[:200]
	return None


def _extract_article1_section(text: str) -> str:
	"""Tenta extrair somente o conteúdo do Art. 1º (ou Artigo 1º).

	Heurística:
	- Busca início por "Art. 1º", "Art. 1o", "Artigo 1º" (case-insensitive)
	- Corta até antes de "Art. 2º/2o/Artigo 2º" ou final do texto
	"""
	import re
	if not text:
		return ""
	t = text
	# normalizar espaços para facilitar o recorte
	t = re.sub(r"\s+", " ", t).strip()

	# localizar início do art. 1º (com ou sem o prefixo "Art." ou "Artigo")
	m1 = re.search(r"\b(?:Art\.?|Artigo)\s*1(º|o)?\b[:\-]?", t, flags=re.I)
	if not m1:
		return ""
	start = m1.start()

	# localizar início do art. 2º a partir do fim do match 1 (com ou sem prefixo)
	rest = t[m1.end():]
	m2 = re.search(r"\b(?:Art\.?|Artigo)\s*2(º|o)?\b", rest, flags=re.I)
	if m2:
		end = m1.end() + m2.start()
	else:
		end = len(t)

	return t[start:end].strip()


def _strip_legalese_preamble(text: str) -> str:
	"""Remove trechos iniciais de formalidades jurídicas e corta até a parte dispositiva.

	Heurísticas:
	- Descarta tudo até (e incluindo) marcadores como "resolve:", "resolvo:", "decide:".
	- Remove cabeçalhos como "O MINISTRO...", "A MINISTRA...", "no uso de suas atribuições".
	- Remove blocos iniciados por "tendo em vista", "considerando", "nos termos", "com fundamento".
	- Normaliza "Art. 1º" para "1º" (remove o token "Art.").
	"""
	import re
	if not text:
		return ""

	t = text
	# Normalizar quebras para facilitar recortes
	t = re.sub(r"\s+", " ", t).strip()

	low = t.lower()
	markers = ["resolve:", "resolvo:", "decide:", "decido:", "torna público:", "torno público:", "torna publico:", "torno publico:"]
	cut_idx = -1
	for m in markers:
		i = low.find(m)
		if i >= 0:
			cut_idx = max(cut_idx, i + len(m))
			break
	if cut_idx >= 0:
		t = t[cut_idx:].lstrip(" -:;—")
		low = t.lower()

	# Remover preâmbulos comuns no início
	preambles = [
		r"^(o|a)\s+minist[roa]\s+de\s+estado.*?\b",  # O MINISTRO DE ESTADO...
		r"no\s+uso\s+de\s+suas\s+atribui[cç][oõ]es.*?\b",
		r"tendo\s+em\s+vista.*?\b",
		r"considerando.*?\b",
		r"nos\s+termos\s+do.*?\b",
		r"com\s+fundamento\s+no.*?\b",
		r"de\s+acordo\s+com.*?\b",
	]
	for pat in preambles:
		t = re.sub(pat, "", t, flags=re.I)
		t = t.strip(" -:;— ")

	# Normalizar "Art." prefixo antes do ordinal
	t = re.sub(r"\bArt\.?\s*(\d+º?)", r"\1", t, flags=re.I)
	return t.strip()


def _first_sentences(text: str, max_sents: int = 2) -> str:
	import re
	sents = [s.strip() for s in re.split(r"[.!?]\s+", text) if s.strip()]
	if not sents:
		return ""
	out = ". ".join(sents[:max_sents])
	return out + ("" if out.endswith(".") else ".")


def _cap_sentences(text: str, max_sents: int) -> str:
	"""Corta o texto para no máximo N frases simples, preservando ponto final."""
	import re
	if max_sents <= 0 or not text:
		return text or ""
	sents = [s.strip() for s in re.split(r"[.!?]\s+", text) if s.strip()]
	if not sents:
		return ""
	out = ". ".join(sents[:max_sents])
	return out + ("" if out.endswith(".") else ".")


def _final_clean_snippet(snippet: str) -> str:
	"""Limpa eventuais resíduos no snippet: metadados, espaços e múltiplos pontos."""
	import re
	if not snippet:
		return ""
	s = _remove_dou_metadata(snippet)
	s = re.sub(r"\s+", " ", s).strip()
	s = re.sub(r"\.\.+", ".", s)
	return s


def _make_bullet_title_from_text(text: str, max_len: int = 140) -> Optional[str]:
	"""Gera um título amigável a partir do texto já limpo de juridiquês."""
	import re
	if not text:
		return None
	head = _first_sentences(text, 1)
	if not head:
		return None
	head = re.sub(r"^\d+º\s+", "", head).strip()  # remove ordinal inicial
	return head[:max_len]


def _default_simple_summarizer(text: str, max_lines: int, mode: str, keywords=None) -> str:
	"""
	Sumarizador simples fallback quando nenhum outro é fornecido.
	Aplica limpeza de preâmbulo jurídico e extrai frases do início ou centro.
	"""
	import re
	clean = _strip_legalese_preamble(text)
	# priorizar somente o Art. 1º, se existir
	a1 = _extract_article1_section(clean)
	base = a1 or clean
	sents = [s.strip() for s in re.split(r"[.!?]\s+", base) if s.strip()]

	if not sents:
		return ""

	if mode in ("head", "lead"):
		result = ". ".join(sents[:max_lines])
		return result + ("" if result.endswith(".") else ".")

	# mode "center"
	mid = max(0, (len(sents) // 2) - (max_lines // 2))
	chunk = sents[mid: mid + max_lines]
	result = ". ".join(chunk)
	return result + ("" if result.endswith(".") else ".")


def _mk_suffix(it: Dict[str, Any]) -> str:
	"""
	Cria um sufixo padronizado com metadados do item (data, seção, edição, página).
    
	Returns:
		String formatada com metadados, ou string vazia se não houver dados
	"""
	parts = []
    
	if it.get("data_publicacao"):
		parts.append(it["data_publicacao"])
        
	if it.get("secao"):
		parts.append(it["secao"])
        
	if it.get("edicao"):
		parts.append(f"Edição {it['edicao']}")
        
	if it.get("pagina"):
		parts.append(f"p. {it['pagina']}")
        
	return (" — " + " • ".join(parts)) if parts else ""


def _minimal_summary_from_item(it: Dict[str, Any]) -> Optional[str]:
	"""Último recurso: construir um resumo mínimo a partir do cabeçalho ou título."""
	try:
		head = _extract_doc_header_line(it)
	except Exception:
		head = None
	if head:
		return _final_clean_snippet(head)
	t = it.get("title_friendly") or it.get("titulo") or it.get("titulo_listagem") or ""
	if t:
		return _final_clean_snippet(str(t))
	return None


def _summarize_item(
	it: Dict[str, Any], 
	summarizer_fn: Optional[Callable], 
	summarize: bool, 
	keywords: Optional[List[str]], 
	max_lines: int, 
	mode: str
) -> Optional[str]:
	"""
	Aplica sumarização a um item se summarize=True e summarizer_fn disponível.
	Lida com diferentes assinaturas de summarizer_fn.
    
	Returns:
		String resumida ou None se não foi possível resumir
	"""
	if not summarize or not summarizer_fn:
		return None
        
	base = it.get("texto") or it.get("ementa") or ""
	if not base:
		return None
	# Separar cabeçalho do corpo para o resumo não repetir o cabeçalho do ato
	use_base = base
	try:
		header, body = _split_doc_header(base)
		# Só usar o body se ele tiver conteúdo razoável (evita ficar vazio quando só há título)
		if body and len(body.strip()) >= 30:
			use_base = body
	except Exception:
		pass
	# Remover metadados do DOU, limpar juridiquês e tentar extrair somente o Art. 1º
	try:
		clean = _remove_dou_metadata(use_base)
		clean = _strip_legalese_preamble(clean)
		a1 = _extract_article1_section(clean)
		base_eff = a1 or clean
		if not base_eff:
			# fallback: tentar com o texto original limpo (sem remover cabeçalho)
			clean2 = _strip_legalese_preamble(_remove_dou_metadata(base))
			base_eff = clean2 or base
		base = base_eff
	except Exception:
		pass
        
	try:
		snippet = summarizer_fn(base, max_lines, mode, keywords)
	except TypeError:
		# Compatibilidade com summarizers que aceitam (text, max_lines, mode) apenas
		try:
			snippet = summarizer_fn(base, max_lines, mode)
		except Exception as e:
			logger.warning(f"Erro ao sumarizar: {e}")
			return None
	except Exception as e:
		logger.warning(f"Erro ao sumarizar: {e}")
		return None
	# Se o summarizer retornar vazio, tentar fallback com a base original
	try:
		if not snippet:
			alt = _strip_legalese_preamble(_remove_dou_metadata(base)) if base else base
			if alt and alt != base:
				try:
					snippet = summarizer_fn(alt, max_lines, mode, keywords)
				except TypeError:
					snippet = summarizer_fn(alt, max_lines, mode)
	except Exception:
		pass

	# Se ainda não houver snippet, aplicar fallback com sumarizador simples padrão
	if not snippet:
		try:
			snippet = _default_simple_summarizer(base, max_lines, mode, keywords)
		except Exception:
			snippet = None

	# Pós-processamento: limitar a N frases e limpar resíduos/metadata
	if snippet:
		snippet = _cap_sentences(snippet, max_lines)
		snippet = _final_clean_snippet(snippet)
		# Salvaguarda: se após o pós-processamento o resumo ficar vazio, reconstruir com header/título
		if not snippet.strip():
			try:
				head2 = _extract_doc_header_line(it)
			except Exception:
				head2 = None
			if head2:
				snippet = _final_clean_snippet(head2)
			else:
				t2 = it.get("title_friendly") or it.get("titulo") or it.get("titulo_listagem") or ""
				if t2:
					snippet = _final_clean_snippet(str(t2))
				else:
					# Último recurso: usar sumarizador simples no texto-base
					try:
						snippet = _default_simple_summarizer(base or "", max_lines, mode, keywords)
						snippet = _final_clean_snippet(snippet)
					except Exception:
						snippet = ""
	return snippet


class BulletinGenerator(ABC):
	"""Classe base abstrata para geradores de boletim em diferentes formatos."""
    
	def __init__(
		self,
		result: Dict[str, Any],
		out_path: str,
		summarizer_fn: Optional[Callable],
		summarize: bool,
		keywords: Optional[List[str]],
		max_lines: int,
		mode: str
	):
		self.result = result
		self.out_path = out_path
		self.summarizer_fn = summarizer_fn
		self.summarize = summarize
		self.keywords = keywords
		self.max_lines = max_lines
		self.mode = mode
		self.date = result.get("data", "")
		self.secao = result.get("secao", "")
        
		# Agrupar itens por (órgão, tipo_ato)
		self.grouped = self._group_items()
        
	def _group_items(self) -> Dict[Tuple[str, str], List[Dict[str, Any]]]:
		"""Agrupa itens por (órgão, tipo_ato)."""
		grouped = defaultdict(list)
		for it in self.result.get("itens", []):
			org = it.get("orgao") or "Sem órgão"
			tipo = it.get("tipo_ato") or "Sem tipo"
			grouped[(org, tipo)].append(it)
		return grouped
        
	def generate(self) -> Dict[str, Any]:
		"""
		Gera o boletim no formato específico e retorna metadados.
        
		Returns:
			Dict com metadados: groups, items, summarized, output
		"""
		# Preparar diretório de saída
		Path(self.out_path).parent.mkdir(parents=True, exist_ok=True)
        
		# Geração específica por formato
		summarized = self._generate_content()
        
		return {
			"groups": len(self.grouped),
			"items": sum(len(v) for v in self.grouped.values()),
			"summarized": summarized,
			"output": self.out_path
		}
    
	@abstractmethod
	def _generate_content(self) -> int:
		"""
		Implementação específica da geração de conteúdo para cada formato.
        
		Returns:
			Número de itens sumarizados
		"""
		pass


class DocxBulletinGenerator(BulletinGenerator):
	"""Gerador de boletim em formato DOCX."""
    
	def _generate_content(self) -> int:
		try:
			from docx import Document
			from docx.oxml import OxmlElement
			from docx.oxml.ns import qn
			from docx.opc.constants import RELATIONSHIP_TYPE as RT
		except ImportError:
			logger.error("Módulo python-docx não encontrado. Instale com: pip install python-docx")
			raise
            
		def add_hyperlink(paragraph, url: str, text: str, color="0000FF", underline=True):
			"""Adiciona hyperlink a um parágrafo no DOCX."""
			try:
				r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
			except Exception:
				return
                
			hyperlink = OxmlElement('w:hyperlink')
			hyperlink.set(qn('r:id'), r_id)
            
			new_run = OxmlElement('w:r')
			rPr = OxmlElement('w:rPr')
            
			if color:
				c = OxmlElement('w:color')
				c.set(qn('w:val'), color)
				rPr.append(c)
                
			if underline:
				u = OxmlElement('w:u')
				u.set(qn('w:val'), 'single')
				rPr.append(u)
                
			new_run.append(rPr)
			t = OxmlElement('w:t')
			t.text = text
			new_run.append(t)
			hyperlink.append(new_run)
			paragraph._p.append(hyperlink)

		# Criar documento e adicionar título
		doc = Document()
		doc.add_heading(f"Boletim DOU — {self.date} ({self.secao})", 0)

		# Contador de itens sumarizados
		summarized = 0
        
		# Para cada grupo (órgão + tipo de ato)
		for (org, tipo), arr in self.grouped.items():
			doc.add_heading(f"{org} — {tipo}", level=1)
            
			# Para cada item no grupo
			for it in arr:
				base_text = it.get("texto") or it.get("ementa") or ""
				cleaned = _strip_legalese_preamble(_remove_dou_metadata(base_text)) if base_text else ""
				# Manter texto do link inalterado (sem derivar título do corpo)
				titulo = it.get("title_friendly") or it.get("titulo") or it.get("titulo_listagem") or "Sem título"
				durl = it.get("detail_url") or it.get("link") or ""
				pdf = it.get("pdf_url") or ""
				suffix = _mk_suffix(it)

				# Adicionar título com links
				p = doc.add_paragraph(style="List Bullet")
				if durl:
					add_hyperlink(p, durl, titulo)
				else:
					p.add_run(titulo)
                    
				if pdf:
					p.add_run(" [")
					add_hyperlink(p, pdf, "PDF")
					p.add_run("]")
                    
				if suffix:
					p.add_run(suffix)

				# Adicionar resumo se disponível
				snippet = _summarize_item(it, self.summarizer_fn, self.summarize, 
										 self.keywords, self.max_lines, self.mode)
				if snippet:
					summarized += 1
					pr = doc.add_paragraph()
					r = pr.add_run("Resumo: ")
					r.bold = True
					pr.add_run(snippet)

		# Salvar documento
		doc.save(self.out_path)
		return summarized


class MarkdownBulletinGenerator(BulletinGenerator):
	"""Gerador de boletim em formato Markdown."""
    
	def _generate_content(self) -> int:
		lines = [f"# Boletim DOU — {self.date} ({self.secao})", ""]
		summarized = 0
        
		for (org, tipo), arr in self.grouped.items():
			lines.append(f"## {org} — {tipo}")
			lines.append("")
            
			for it in arr:
				base_text = it.get("texto") or it.get("ementa") or ""
				cleaned = _strip_legalese_preamble(_remove_dou_metadata(base_text)) if base_text else ""
				# Manter texto do link inalterado
				titulo = it.get("title_friendly") or it.get("titulo") or it.get("titulo_listagem") or "Sem título"
				durl = it.get("detail_url") or it.get("link") or ""
				pdf = it.get("pdf_url") or ""
				suffix = _mk_suffix(it)
                
				# Link markdown para o título
				base_line = f"- [{titulo}]({durl})" if durl else f"- {titulo}"
                
				# Adicionar link para PDF se disponível
				if pdf:
					base_line += f" [PDF]({pdf})"
                    
				if suffix:
					base_line += suffix
                    
				lines.append(base_line)
                
				# Adicionar resumo se disponível
				snippet = _summarize_item(it, self.summarizer_fn, self.summarize, 
								 self.keywords, self.max_lines, self.mode)
				if not snippet and self.summarize:
					snippet = _minimal_summary_from_item(it)
				if snippet:
					summarized += 1
					lines.append(f"  \n  _Resumo:_ {snippet}")
                    
				lines.append("")
                
		# Escrever arquivo markdown
		Path(self.out_path).write_text("\n".join(lines), encoding="utf-8")
		return summarized


class HtmlBulletinGenerator(BulletinGenerator):
	"""Gerador de boletim em formato HTML."""
    
	def _generate_content(self) -> int:
		parts = [f"<h1>Boletim DOU — {html_lib.escape(self.date)} ({html_lib.escape(self.secao)})</h1>"]
		summarized = 0
        
		for (org, tipo), arr in self.grouped.items():
			parts.append(f"<h2>{html_lib.escape(org)} — {html_lib.escape(tipo)}</h2>")
			parts.append("<ul>")
            
			for it in arr:
				base_text = it.get("texto") or it.get("ementa") or ""
				cleaned = _strip_legalese_preamble(_remove_dou_metadata(base_text)) if base_text else ""
				# Manter texto do link inalterado
				titulo = it.get("title_friendly") or it.get("titulo") or it.get("titulo_listagem") or "Sem título"
				durl = it.get("detail_url") or it.get("link") or ""
				pdf = it.get("pdf_url") or ""
				suffix = _mk_suffix(it)
                
				# Link HTML para título
				title_html = html_lib.escape(titulo)
				if durl:
					title_html = f'<a href="{html_lib.escape(durl)}">{title_html}</a>'
                    
				# Link para PDF
				pdf_html = f' <a href="{html_lib.escape(pdf)}">[PDF]</a>' if pdf else ""
                
				parts.append(f"<li>{title_html}{pdf_html}{html_lib.escape(suffix)}")
                
				# Adicionar resumo se disponível
				snippet = _summarize_item(it, self.summarizer_fn, self.summarize, 
								 self.keywords, self.max_lines, self.mode)
				if not snippet and self.summarize:
					snippet = _minimal_summary_from_item(it)
				if snippet:
					summarized += 1
					parts.append(f"<div><strong>Resumo:</strong> {html_lib.escape(snippet)}</div>")
                    
				parts.append("</li>")
                
			parts.append("</ul>")
            
		# Escrever arquivo HTML
		Path(self.out_path).write_text("\n".join(parts), encoding="utf-8")
		return summarized


def generate_bulletin(
	result: Dict[str, Any],
	out_path: str,
	kind: str = "docx",
	summarize: bool = False,
	summarizer: Optional[Callable[[str, int, str, Optional[List[str]]], str]] = None,
	keywords: Optional[List[str]] = None,
	max_lines: int = 5,
	mode: str = "center"
) -> Dict[str, Any]:
	"""
	Gera boletim e retorna metadados.
    
	Args:
		result: Dicionário com dados do resultado (data, secao, itens)
		out_path: Caminho para arquivo de saída
		kind: Formato do boletim (docx, md, html)
		summarize: Se True, inclui resumo para cada item
		summarizer: Função de sumarização personalizada
		keywords: Lista de palavras-chave para sumarização
		max_lines: Número máximo de linhas no resumo
		mode: Modo de sumarização (center, head)
        
	Returns:
		Dict com metadados: {groups, items, summarized, output}
	"""
	# Definir summarizer_fn
	summarizer_fn = summarizer
	if summarize and summarizer_fn is None:
		summarizer_fn = _default_simple_summarizer  # fallback
    
	# Criar gerador apropriado conforme formato
	if kind == "docx":
		generator = DocxBulletinGenerator(
			result, out_path, summarizer_fn, summarize, keywords, max_lines, mode
		)
	elif kind == "md":
		generator = MarkdownBulletinGenerator(
			result, out_path, summarizer_fn, summarize, keywords, max_lines, mode
		)
	elif kind == "html":
		generator = HtmlBulletinGenerator(
			result, out_path, summarizer_fn, summarize, keywords, max_lines, mode
		)
	else:
		raise ValueError(f"Formato '{kind}' não suportado. Use: docx|md|html")
    
	# Gerar e retornar metadados
	return generator.generate()
