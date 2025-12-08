"""
bulletin_utils.py
Geração de boletins em DOCX, Markdown e HTML com agrupamento por (órgão, tipo_ato)
e sumarização (simples ou avançada) opcional.

Função principal:
  generate_bulletin(result_dict, out_path, kind="docx", summarize=False,
                    summarizer=None, keywords=None, max_lines=5, mode="center")
"""

from __future__ import annotations

import contextlib
import html as html_lib
import re
from abc import ABC, abstractmethod
from collections import defaultdict
from collections.abc import Callable
from pathlib import Path
from typing import Any

from .log_utils import get_logger
from .text_cleaning import (
    cap_sentences as _cap_sentences,
    extract_article1_section as _extract_article1_section,
    extract_doc_header_line as _extract_doc_header_line,
    final_clean_snippet as _final_clean_snippet,
    remove_dou_metadata as _remove_dou_metadata,
    split_doc_header as _split_doc_header,
    strip_legalese_preamble as _strip_legalese_preamble,
)

logger = get_logger(__name__)


def _default_simple_summarizer(text: str, max_lines: int, mode: str, keywords=None) -> str:
    """
    Sumarizador simples fallback quando nenhum outro é fornecido.
    Aplica limpeza de preâmbulo jurídico e extrai frases do início ou centro.
    """
    clean = _strip_legalese_preamble(text)
    # priorizar somente o Art. 1º, se existir
    a1 = _extract_article1_section(clean)
    base = a1 or clean
    sents = [s.strip() for s in re.split(r"[.!?]\s+", base) if s.strip()]

    if not sents:
        # Fallback: se não houver pontuação, sintetizar com primeiras palavras
        words = re.findall(r"\w+[\w-]*", base)
        if not words:
            return ""
        chunk = " ".join(words[: max(5, max_lines * 8)])
        return chunk.strip() + "."

    if mode in ("head", "lead"):
        result = ". ".join(sents[:max_lines])
        return result + ("" if result.endswith(".") else ".")

    # mode "center"
    mid = max(0, (len(sents) // 2) - (max_lines // 2))
    chunk = sents[mid: mid + max_lines]
    result = ". ".join(chunk)
    return result + ("" if result.endswith(".") else ".")


def _mk_suffix(it: dict[str, Any]) -> str:
    """
    Cria um sufixo padronizado com metadados do item (data, seção, edição, página).

    Returns:
        String formatada com metadados, ou string vazia se não houver dados
    """
    parts = []

    if it.get("data_publicacao"):
        parts.append(it["data_publicacao"])

    if it.get("secao"):
        parts.append(it["secao"])

    if it.get("edicao"):
        parts.append(f"Edição {it['edicao']}")

    if it.get("pagina"):
        parts.append(f"p. {it['pagina']}")

    return (" — " + " • ".join(parts)) if parts else ""


def _minimal_summary_from_item(it: dict[str, Any]) -> str | None:
    """Último recurso para obter um resumo mínimo a partir do cabeçalho ou título.

    Retorna uma string curta e limpa ou None.
    """
    try:
        head = _extract_doc_header_line(it)
    except Exception as e:
        logger.debug(f"Failed to extract doc header: {e}")
        head = None
    if head:
        return _final_clean_snippet(head)
    t = it.get("title_friendly") or it.get("titulo") or it.get("titulo_listagem") or ""
    if t:
        return _final_clean_snippet(str(t))
    return None


def _summarize_item(
    it: dict[str, Any],
    summarizer_fn: Callable | None,
    summarize: bool,
    keywords: list[str] | None,
    max_lines: int,
    mode: str
) -> str | None:
    """
    Aplica sumarização a um item se summarize=True e summarizer_fn disponível.
    Lida com diferentes assinaturas de summarizer_fn.

    Returns:
        String resumida ou None se não foi possível resumir
    """
    from dou_utils.summarization_helpers import (
        extract_base_text,
        get_fallback_from_title,
        prepare_text_for_summarization,
        derive_mode_from_doc_type,
        apply_summarizer_with_fallbacks,
        apply_default_summarizer,
        post_process_snippet,
    )

    if not summarize or not summarizer_fn:
        return None

    # Extract base text
    base = extract_base_text(it)
    if not base:
        return get_fallback_from_title(it)

    # Prepare text: split header/body
    use_base = base
    try:
        _header, body = _split_doc_header(base)
        if body and len(body.strip()) >= 30:
            use_base = body
    except Exception as e:
        logger.debug(f"Failed to split doc header: {e}")

    # Clean and prepare text
    prepared_text = prepare_text_for_summarization(use_base)

    # Derive mode from document type
    derived_mode = derive_mode_from_doc_type(it, mode)

    # Try main summarizer with fallbacks
    snippet, method_tag = apply_summarizer_with_fallbacks(
        summarizer_fn, prepared_text, max_lines, derived_mode, keywords
    )

    # If still no snippet, try default summarizer
    if not snippet:
        snippet = apply_default_summarizer(prepared_text, max_lines, derived_mode, keywords)
        method_tag = method_tag or "default_simple"

    # Final fallback: use header or title
    if not snippet:
        snippet = get_fallback_from_title(it)
        if snippet:
            method_tag = method_tag or "title_fallback"

    # Post-process snippet
    if snippet:
        snippet = post_process_snippet(snippet, max_lines)

        # Safety check: if empty after post-processing, rebuild from title
        if not snippet.strip():
            snippet = get_fallback_from_title(it)
            if snippet:
                snippet = _final_clean_snippet(snippet)
                method_tag = method_tag or "header_line_cleanup"

    return snippet


class BulletinGenerator(ABC):
    """Classe base abstrata para geradores de boletim em diferentes formatos."""

    def __init__(
        self,
        result: dict[str, Any],
        out_path: str,
        summarizer_fn: Callable | None,
        summarize: bool,
        keywords: list[str] | None,
        max_lines: int,
        mode: str
    ):
        self.result = result
        self.out_path = out_path
        self.summarizer_fn = summarizer_fn
        self.summarize = summarize
        self.keywords = keywords
        self.max_lines = max_lines
        self.mode = mode
        self.date = result.get("data", "")
        self.secao = result.get("secao", "")

        # Agrupar itens por (órgão, tipo_ato)
        self.grouped = self._group_items()

    def _group_items(self) -> dict[tuple[str, str], list[dict[str, Any]]]:
        """Agrupa itens por (órgão, tipo_ato)."""
        grouped = defaultdict(list)
        for it in self.result.get("itens", []):
            org = it.get("orgao") or "Sem órgão"
            tipo = it.get("tipo_ato") or "Sem tipo"
            grouped[(org, tipo)].append(it)
        return grouped

    def generate(self) -> dict[str, Any]:
        """
        Gera o boletim no formato específico e retorna metadados.

        Returns:
            Dict com metadados: groups, items, summarized, output
        """
        # Preparar diretório de saída
        Path(self.out_path).parent.mkdir(parents=True, exist_ok=True)

        # Geração específica por formato
        summarized = self._generate_content()

        return {
            "groups": len(self.grouped),
            "items": sum(len(v) for v in self.grouped.values()),
            "summarized": summarized,
            "output": self.out_path
        }

    @abstractmethod
    def _generate_content(self) -> int:
        """
        Implementação específica da geração de conteúdo para cada formato.

        Returns:
            Número de itens sumarizados
        """


class DocxBulletinGenerator(BulletinGenerator):
    """Gerador de boletim em formato DOCX."""

    def _generate_content(self) -> int:
        try:
            from docx import Document
            from docx.opc.constants import RELATIONSHIP_TYPE as RT
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
        except ImportError:
            logger.error("Modulo python-docx nao encontrado. Instale com: pip install python-docx")
            raise

        def add_hyperlink(paragraph, url: str, text: str, color="0000FF", underline=True):
            """Adiciona hyperlink a um paragrafo no DOCX."""
            try:
                r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
            except Exception:
                return

            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)

            new_run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')

            if color:
                c = OxmlElement('w:color')
                c.set(qn('w:val'), color)
                rPr.append(c)

            if underline:
                u = OxmlElement('w:u')
                u.set(qn('w:val'), 'single')
                rPr.append(u)

            new_run.append(rPr)
            t = OxmlElement('w:t')
            t.text = text
            new_run.append(t)
            hyperlink.append(new_run)
            paragraph._p.append(hyperlink)

        def keep_with_next(paragraph):
            """Configura paragrafo para ficar junto com o proximo (evita quebra de pagina)."""
            pPr = paragraph._p.get_or_add_pPr()
            keepNext = OxmlElement('w:keepNext')
            pPr.append(keepNext)

        def keep_together(paragraph):
            """Configura paragrafo para manter todas as linhas juntas."""
            pPr = paragraph._p.get_or_add_pPr()
            keepLines = OxmlElement('w:keepLines')
            pPr.append(keepLines)

        # Criar documento e adicionar titulo
        doc = Document()
        doc.add_heading(f"Boletim DOU - {self.date} ({self.secao})", 0)

        # Contador de itens sumarizados
        summarized = 0

        # Para cada grupo (orgao + tipo de ato)
        for (org, tipo), arr in self.grouped.items():
            doc.add_heading(f"{org} - {tipo}", level=1)

            # Para cada item no grupo
            for it in arr:
                base_text = it.get("texto") or it.get("ementa") or ""
                _strip_legalese_preamble(_remove_dou_metadata(base_text)) if base_text else ""
                # Manter texto do link inalterado (sem derivar titulo do corpo)
                titulo = it.get("title_friendly") or it.get("titulo") or it.get("titulo_listagem") or "Sem titulo"
                durl = it.get("detail_url") or it.get("link") or ""
                pdf = it.get("pdf_url") or ""
                suffix = _mk_suffix(it)

                # Adicionar titulo com links
                p = doc.add_paragraph(style="List Bullet")
                if durl:
                    add_hyperlink(p, durl, titulo)
                else:
                    p.add_run(titulo)

                if pdf:
                    p.add_run(" [")
                    add_hyperlink(p, pdf, "PDF")
                    p.add_run("]")

                if suffix:
                    p.add_run(suffix)

                # Adicionar resumo se disponivel
                snippet = _summarize_item(it, self.summarizer_fn, self.summarize,
                                         self.keywords, self.max_lines, self.mode)
                if not snippet and self.summarize:
                    snippet = _minimal_summary_from_item(it)
                if snippet:
                    summarized += 1
                    # Manter titulo junto com resumo (evita quebra de pagina entre eles)
                    keep_with_next(p)
                    pr = doc.add_paragraph()
                    keep_together(pr)  # Manter linhas do resumo juntas
                    r = pr.add_run("Resumo: ")
                    r.bold = True
                    pr.add_run(snippet)

        # Salvar documento
        doc.save(self.out_path)
        return summarized


class MarkdownBulletinGenerator(BulletinGenerator):
    """Gerador de boletim em formato Markdown."""

    def _generate_content(self) -> int:
        lines = [f"# Boletim DOU — {self.date} ({self.secao})", ""]
        summarized = 0

        for (org, tipo), arr in self.grouped.items():
            lines.append(f"## {org} — {tipo}")
            lines.append("")

            for it in arr:
                base_text = it.get("texto") or it.get("ementa") or ""
                _strip_legalese_preamble(_remove_dou_metadata(base_text)) if base_text else ""
                # Manter texto do link inalterado
                titulo = it.get("title_friendly") or it.get("titulo") or it.get("titulo_listagem") or "Sem título"
                durl = it.get("detail_url") or it.get("link") or ""
                pdf = it.get("pdf_url") or ""
                suffix = _mk_suffix(it)

                # Link markdown para o título
                base_line = f"- [{titulo}]({durl})" if durl else f"- {titulo}"

                # Adicionar link para PDF se disponível
                if pdf:
                    base_line += f" [PDF]({pdf})"

                if suffix:
                    base_line += suffix

                lines.append(base_line)

                # Adicionar resumo se disponível
                snippet = _summarize_item(it, self.summarizer_fn, self.summarize,
                                         self.keywords, self.max_lines, self.mode)
                if not snippet and self.summarize:
                    snippet = _minimal_summary_from_item(it)
                if snippet:
                    summarized += 1
                    lines.append(f"  \n  _Resumo:_ {snippet}")

                lines.append("")

        # Escrever arquivo markdown
        Path(self.out_path).write_text("\n".join(lines), encoding="utf-8")
        return summarized


class HtmlBulletinGenerator(BulletinGenerator):
    """Gerador de boletim em formato HTML."""

    def _generate_content(self) -> int:
        parts = [f"<h1>Boletim DOU — {html_lib.escape(self.date)} ({html_lib.escape(self.secao)})</h1>"]
        summarized = 0

        for (org, tipo), arr in self.grouped.items():
            parts.append(f"<h2>{html_lib.escape(org)} — {html_lib.escape(tipo)}</h2>")
            parts.append("<ul>")

            for it in arr:
                base_text = it.get("texto") or it.get("ementa") or ""
                _strip_legalese_preamble(_remove_dou_metadata(base_text)) if base_text else ""
                # Manter texto do link inalterado
                titulo = it.get("title_friendly") or it.get("titulo") or it.get("titulo_listagem") or "Sem título"
                durl = it.get("detail_url") or it.get("link") or ""
                pdf = it.get("pdf_url") or ""
                suffix = _mk_suffix(it)

                # Link HTML para título
                title_html = html_lib.escape(titulo)
                if durl:
                    title_html = f'<a href="{html_lib.escape(durl)}">{title_html}</a>'

                # Link para PDF
                pdf_html = f' <a href="{html_lib.escape(pdf)}">[PDF]</a>' if pdf else ""

                parts.append(f"<li>{title_html}{pdf_html}{html_lib.escape(suffix)}")

                # Adicionar resumo se disponível
                snippet = _summarize_item(it, self.summarizer_fn, self.summarize,
                                         self.keywords, self.max_lines, self.mode)
                if not snippet and self.summarize:
                    snippet = _minimal_summary_from_item(it)
                if snippet:
                    summarized += 1
                    parts.append(f"<div><strong>Resumo:</strong> {html_lib.escape(snippet)}</div>")

                parts.append("</li>")

            parts.append("</ul>")

        # Escrever arquivo HTML
        Path(self.out_path).write_text("\n".join(parts), encoding="utf-8")
        return summarized


def generate_bulletin(
    result: dict[str, Any],
    out_path: str,
    kind: str = "docx",
    summarize: bool = False,
    summarizer: Callable[[str, int, str, list[str] | None], str] | None = None,
    keywords: list[str] | None = None,
    max_lines: int = 5,
    mode: str = "center"
) -> dict[str, Any]:
    """
    Gera boletim e retorna metadados.

    Args:
        result: Dicionário com dados do resultado (data, secao, itens)
        out_path: Caminho para arquivo de saída
        kind: Formato do boletim (docx, md, html)
        summarize: Se True, inclui resumo para cada item
        summarizer: Função de sumarização personalizada
        keywords: Lista de palavras-chave para sumarização
        max_lines: Número máximo de linhas no resumo
        mode: Modo de sumarização (center, head)

    Returns:
        Dict com metadados: {groups, items, summarized, output}
    """
    # Definir summarizer_fn
    summarizer_fn = summarizer
    if summarize and summarizer_fn is None:
        summarizer_fn = _default_simple_summarizer  # fallback

    # Criar gerador apropriado conforme formato
    if kind == "docx":
        generator = DocxBulletinGenerator(
            result, out_path, summarizer_fn, summarize, keywords, max_lines, mode
        )
    elif kind == "md":
        generator = MarkdownBulletinGenerator(
            result, out_path, summarizer_fn, summarize, keywords, max_lines, mode
        )
    elif kind == "html":
        generator = HtmlBulletinGenerator(
            result, out_path, summarizer_fn, summarize, keywords, max_lines, mode
        )
    else:
        raise ValueError(f"Formato '{kind}' não suportado. Use: docx|md|html")

    # Gerar e retornar metadados
    return generator.generate()
