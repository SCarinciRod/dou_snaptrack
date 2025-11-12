"""
eagendas_document.py
Geração de documentos DOCX para E-Agendas com agrupamento por agente.

Estrutura do documento:
- Organizado por Agente (não por órgão como no DOU)
- Cada agente tem seus compromissos listados por data
- Formato informal/simples (sem regex ou limpeza agressiva)
- Metadados: Órgão, Cargo, Período

Função principal:
  generate_eagendas_document(events_data, out_path, include_metadata=True)
"""

from __future__ import annotations

from collections import defaultdict
from datetime import datetime
from pathlib import Path
from typing import Any

try:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Pt, RGBColor
except ImportError:
    Document = None
    WD_PARAGRAPH_ALIGNMENT = None
    Pt = None
    RGBColor = None

from .log_utils import get_logger

logger = get_logger(__name__)


def _format_date(date_str: str) -> str:
    """
    Formata data de YYYY-MM-DD para formato legível (DD/MM/YYYY).

    Args:
        date_str: Data em formato ISO (YYYY-MM-DD)

    Returns:
        Data formatada ou original se falhar
    """
    try:
        dt = datetime.strptime(date_str, "%Y-%m-%d")
        return dt.strftime("%d/%m/%Y")
    except Exception:
        return date_str


def _group_events_by_agent(events_data: dict[str, Any]) -> dict[str, list[dict[str, Any]]]:
    """
    Agrupa eventos por agente público.

    Args:
        events_data: Estrutura JSON com eventos coletados
            Formato esperado:
            {
                "periodo": {"inicio": "YYYY-MM-DD", "fim": "YYYY-MM-DD"},
                "agentes": [
                    {
                        "orgao": {"id": "X", "nome": "Nome Órgão"},
                        "cargo": {"id": "Y", "nome": "Nome Cargo"},
                        "agente": {"id": "Z", "nome": "Nome Agente"},
                        "eventos": {
                            "YYYY-MM-DD": [
                                {"title": "...", "time": "...", "type": "...", "details": "..."},
                                ...
                            ],
                            ...
                        }
                    },
                    ...
                ]
            }

    Returns:
        Dict: {"Nome Agente": [evento1, evento2, ...], ...}
    """
    grouped = defaultdict(list)

    try:
        agentes_list = events_data.get("agentes", [])
        for agente_data in agentes_list:
            agente_nome = agente_data.get("agente", {}).get("nome", "Agente Desconhecido")
            orgao_nome = agente_data.get("orgao", {}).get("nome", "")
            cargo_nome = agente_data.get("cargo", {}).get("nome", "")

            eventos_por_dia = agente_data.get("eventos", {})

            for date_str, eventos_list in eventos_por_dia.items():
                for evento in eventos_list:
                    # Adicionar metadados de contexto
                    evento_full = {
                        **evento,
                        "date": date_str,
                        "agente_nome": agente_nome,
                        "orgao_nome": orgao_nome,
                        "cargo_nome": cargo_nome,
                    }
                    grouped[agente_nome].append(evento_full)

        # Ordenar eventos de cada agente por data
        for agente_nome in grouped:
            grouped[agente_nome].sort(key=lambda e: e.get("date", ""))

        return dict(grouped)

    except Exception as e:
        logger.error(f"Erro ao agrupar eventos por agente: {e}")
        return {}


def _add_header(doc, title: str, subtitle: str | None = None):
    """
    Adiciona cabeçalho ao documento.

    Args:
        doc: Documento python-docx
        title: Título principal
        subtitle: Subtítulo opcional
    """
    if not doc or WD_PARAGRAPH_ALIGNMENT is None or Pt is None or RGBColor is None:
        return

    # Título principal
    h = doc.add_heading(title, level=0)
    h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    h.runs[0].font.color.rgb = RGBColor(0, 51, 102)  # Azul escuro

    # Subtítulo
    if subtitle:
        p = doc.add_paragraph(subtitle)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.runs[0].font.size = Pt(11)
        p.runs[0].italic = True

    # Espaço
    doc.add_paragraph()


def _add_agent_section(doc, agente_nome: str, eventos: list[dict[str, Any]]):
    """
    Adiciona seção de um agente com seus compromissos.

    Args:
        doc: Documento python-docx
        agente_nome: Nome do agente público
        eventos: Lista de eventos do agente
    """
    if not doc or not eventos or Pt is None or RGBColor is None:
        return

    # Cabeçalho do agente
    h = doc.add_heading(agente_nome, level=1)
    h.runs[0].font.color.rgb = RGBColor(0, 102, 204)  # Azul médio

    # Metadados do agente (órgão e cargo do primeiro evento)
    primeiro = eventos[0]
    orgao = primeiro.get("orgao_nome", "")
    cargo = primeiro.get("cargo_nome", "")

    if orgao or cargo:
        meta_parts = []
        if orgao:
            meta_parts.append(f"Órgão: {orgao}")
        if cargo:
            meta_parts.append(f"Cargo: {cargo}")
        p = doc.add_paragraph(" | ".join(meta_parts))
        p.runs[0].font.size = Pt(9)
        p.runs[0].italic = True
        p.runs[0].font.color.rgb = RGBColor(102, 102, 102)  # Cinza

    doc.add_paragraph()  # Espaço

    # Agrupar eventos por data
    eventos_por_data = defaultdict(list)
    for evento in eventos:
        date_str = evento.get("date", "")
        eventos_por_data[date_str].append(evento)

    # Listar eventos por data
    for date_str in sorted(eventos_por_data.keys()):
        # Cabeçalho da data
        date_formatted = _format_date(date_str)
        h_date = doc.add_heading(f"📅 {date_formatted}", level=2)
        h_date.runs[0].font.color.rgb = RGBColor(51, 51, 51)  # Cinza escuro

        # Eventos deste dia
        eventos_dia = eventos_por_data[date_str]
        for evento in eventos_dia:
            title = evento.get("title", "Compromisso sem título")
            time_str = evento.get("time", "")
            tipo = evento.get("type", "")
            details = evento.get("details", "")

            # Linha do evento
            p = doc.add_paragraph(style="List Number")
            p.add_run(title).bold = True

            # Horário (se houver)
            if time_str:
                p.add_run(f"\n⏰ {time_str}")

            # Tipo (se houver)
            if tipo:
                p.add_run(f"\n🏷️  {tipo}")

            # Detalhes (se houver)
            if details:
                p.add_run(f"\n📝 {details}")

            p.runs[0].font.size = Pt(10)

        doc.add_paragraph()  # Espaço entre datas

    # Separador entre agentes
    doc.add_paragraph("_" * 80)
    doc.add_paragraph()


def generate_eagendas_document(
    events_data: dict[str, Any],
    out_path: str | Path,
    include_metadata: bool = True,
    title: str | None = None
) -> dict[str, Any]:
    """
    Gera documento DOCX com agendas dos agentes públicos.

    Args:
        events_data: Estrutura JSON com eventos coletados (veja _group_events_by_agent)
        out_path: Caminho para arquivo de saída (.docx)
        include_metadata: Se True, inclui metadados (órgão, cargo) em cada agente
        title: Título personalizado do documento (opcional)

    Returns:
        Dict com metadados: {
            "agents": int,        # Número de agentes
            "events": int,        # Total de eventos
            "period": str,        # Período formatado
            "output": str         # Caminho do arquivo gerado
        }

    Raises:
        ImportError: Se python-docx não estiver instalado
        Exception: Se houver erro na geração do documento
    """
    if Document is None:
        raise ImportError(
            "python-docx não está instalado. Execute: pip install python-docx"
        )

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    try:
        # Agrupar eventos por agente
        events_by_agent = _group_events_by_agent(events_data)
        total_events = sum(len(eventos) for eventos in events_by_agent.values())
        num_agents = len(events_by_agent)

        logger.info(f"Gerando documento para {num_agents} agentes com {total_events} eventos")

        # Criar documento
        doc = Document()

        # Período
        periodo = events_data.get("periodo", {})
        inicio = _format_date(periodo.get("inicio", ""))
        fim = _format_date(periodo.get("fim", ""))
        periodo_str = f"Período: {inicio} a {fim}" if inicio and fim else ""

        # Cabeçalho
        doc_title = title or "Agendas de Agentes Públicos"
        _add_header(doc, doc_title, periodo_str)

        # Processar cada agente (ordenado alfabeticamente)
        for agente_nome in sorted(events_by_agent.keys()):
            eventos = events_by_agent[agente_nome]
            _add_agent_section(doc, agente_nome, eventos)

        # Salvar documento
        doc.save(str(out_path))
        logger.info(f"✅ Documento gerado: {out_path}")

        return {
            "agents": num_agents,
            "events": total_events,
            "period": periodo_str,
            "output": str(out_path.absolute()),
        }

    except Exception as e:
        logger.error(f"Erro ao gerar documento: {e}")
        raise


def generate_eagendas_document_from_json(
    json_path: str | Path,
    out_path: str | Path | None = None,
    include_metadata: bool = True,
    title: str | None = None
) -> dict[str, Any]:
    """
    Gera documento DOCX a partir de arquivo JSON com eventos coletados.

    Args:
        json_path: Caminho para arquivo JSON com eventos
        out_path: Caminho para arquivo de saída (se None, usa mesmo nome que JSON)
        include_metadata: Se True, inclui metadados (órgão, cargo) em cada agente
        title: Título personalizado do documento (opcional)

    Returns:
        Dict com metadados (veja generate_eagendas_document)
    """
    import json

    json_path = Path(json_path)

    if not json_path.exists():
        raise FileNotFoundError(f"Arquivo não encontrado: {json_path}")

    # Carregar dados
    with open(json_path, encoding="utf-8") as f:
        events_data = json.load(f)

    # Determinar caminho de saída
    if out_path is None:
        out_path = json_path.with_suffix(".docx")
    else:
        out_path = Path(out_path)

    # Gerar documento
    return generate_eagendas_document(events_data, out_path, include_metadata, title)
