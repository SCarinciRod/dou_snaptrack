# 05_cascade_cli.py (v6.1 - Pipeline 00 -> 05 Plan/Batch)
# Cascata DOU + Coleta + Detalhamento + Batch/Presets + Repetidor + Boletim
# Melhorias: seleção por rótulo, scroll robusto, captura edicao/pagina, faixa de datas,
# dedup (state-file), boletim HTML/MD, e novo modo PLAN para gerar batch a partir do mapa (00).


from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.shared import OxmlElement, qn
from docx.enum.text import WD_UNDERLINE
import argparse
import json
import re
import sys
import unicodedata
import hashlib
import os
import time
from datetime import datetime, timedelta
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from urllib.parse import urljoin, urlparse

from playwright.sync_api import sync_playwright, TimeoutError as PWTimeout

import warnings
warnings.filterwarnings("ignore", category=SyntaxWarning)

# ===== Seletores padrões usados ao mapear/abrir dropdowns =====
DROPDOWN_ROOT_SELECTORS = [
    "[role=combobox]",
    "select",
    "[aria-haspopup=listbox]",
    "[aria-expanded][role=button]",
    "div[class*=select]",
    "div[class*=dropdown]",
    "div[class*=combobox]",
]

LISTBOX_SELECTORS = [
    "[role=listbox]",
    "ul[role=listbox]",
    "div[role=listbox]",
    "ul[role=menu]",
    "div[role=menu]",
    ".ng-dropdown-panel",
    ".p-dropdown-items",
    ".select2-results__options",
    ".rc-virtual-list",
]

OPTION_SELECTORS = [
    "[role=option]",
    "li[role=option]",
    ".ng-option",
    ".p-dropdown-item",
    ".select2-results__option",
    "[data-value]",
    "[data-index]",
]

def _setup_summary_globals(args):
    """
    Lê parâmetros de resumo e popula globais:
      - --summary-lines (preferido)
      - --summary-sentences (fallback)
      - --summary-mode (center|lead|keywords-first)
      - --summary-keywords / --summary-keywords-file
    """
    from pathlib import Path
    global SUMMARY_SENTENCES, SUMMARY_KEYWORDS, SUMMARY_LINES, SUMMARY_MODE

    # Defaults
    SUMMARY_LINES = int(getattr(args, "summary_lines", 0) or 0)  # preferido
    if SUMMARY_LINES <= 0:
        try:
            SUMMARY_SENTENCES = int(getattr(args, "summary_sentences", 7) or 7)
        except Exception:
            SUMMARY_SENTENCES = 7
        SUMMARY_LINES = SUMMARY_SENTENCES
    else:
        SUMMARY_SENTENCES = SUMMARY_LINES

    SUMMARY_MODE = getattr(args, "summary_mode", "center") or "center"
    if SUMMARY_MODE not in ("center", "lead", "keywords-first"):
        SUMMARY_MODE = "center"

    kws = []
    raw = getattr(args, "summary_keywords", None)
    if raw:
        for part in raw.split(";"):
            s = (part or "").strip()
            if s:
                kws.append(s.lower())

    fn = getattr(args, "summary_keywords_file", None)
    if fn:
        try:
            txt = Path(fn).read_text(encoding="utf-8")
            for line in txt.splitlines():
                s = (line or "").strip()
                if s:
                    kws.append(s.lower())
        except Exception:
            pass

    SUMMARY_KEYWORDS = kws

def _apply_summary_overrides_from_job(job: dict):
    """
    Aplica overrides de resumo vindos do job (nível de assunto/topic), 
    prevalecendo sobre os SUMMARY_* globais definidos por _setup_summary_globals(args).
    Aceita:
      - job["summary_lines"] (int > 0)  -> SUMMARY_LINES/SUMMARY_SENTENCES
      - job["summary_mode"]  ("center"|"lead"|"keywords-first") -> SUMMARY_MODE
      - job["summary_keywords"] (list[str]) -> SUMMARY_KEYWORDS (normalizado para lower)
    """
    global SUMMARY_SENTENCES, SUMMARY_KEYWORDS, SUMMARY_LINES, SUMMARY_MODE
    if not isinstance(job, dict):
        return

    # Lines / Sentences
    sl = job.get("summary_lines")
    if isinstance(sl, int) and sl > 0:
        SUMMARY_LINES = sl
        SUMMARY_SENTENCES = sl

    # Mode
    sm = job.get("summary_mode")
    if sm in ("center", "lead", "keywords-first"):
        SUMMARY_MODE = sm

    # Keywords
    kws = job.get("summary_keywords")
    if isinstance(kws, list) and kws:
        SUMMARY_KEYWORDS = [str(k).strip().lower() for k in kws if str(k).strip()]
    

# ------------------------- Utilitários -------------------------
def fmt_date(date_str: Optional[str] = None) -> str:
    if date_str:
        return datetime.strptime(date_str, "%d-%m-%Y").strftime("%d-%m-%Y")
    return datetime.now().strftime("%d-%m-%Y")

def close_cookies(page) -> None:
    for texto in ["ACEITO", "ACEITAR", "OK", "ENTENDI", "CONCORDO", "FECHAR", "ACEITO TODOS"]:
        try:
            btn = page.get_by_role("button", name=re.compile(texto, re.I))
            if btn.count() > 0 and btn.first.is_visible():
                btn.first.click(timeout=1500)
                page.wait_for_timeout(150)
        except Exception:
            pass

def goto(page, data: str, secao: str) -> None:
    # Monta URL de forma determinística (sem depender de BASE global)
    url = f"https://www.in.gov.br/leiturajornal?data={data}&secao={secao}"
    print(f"\n[Abrindo] {url}")
    page.goto(url, wait_until="domcontentloaded", timeout=90_000)
    page.wait_for_load_state("networkidle", timeout=90_000)
    close_cookies(page)
    # Tentar "Visualizar em Lista" OU "Visualizar em Sumário" (regex correta)
    try:
        btn = page.get_by_role("button", name=re.compile(r"(lista|sum[aá]rio)", re.I))
        if btn.count() > 0 and btn.first.is_visible():
            btn.first.click()
            page.wait_for_load_state("networkidle", timeout=60_000)
    except Exception:
        pass

def normalize_text(s: str) -> str:
    if s is None:
        return ""
    s = unicodedata.normalize("NFKD", s)
    s = s.encode("ascii", "ignore").decode("ascii")
    return re.sub(r"\s+", " ", s.strip().lower())

def _best_key_for_option(opt: Dict[str, Any]) -> Tuple[str, Optional[str]]:
    """
    Decide a chave mais estável da opção para selecionar no site.
    Preferência:
      1) value
      2) dataValue
      3) id
      4) dataId
      5) dataIndex
      6) text
    Retorna (key_type, key_value).
    """
    if opt.get("value") not in (None, ""):
        return ("value", str(opt["value"]))
    if opt.get("dataValue") not in (None, ""):
        return ("dataValue", str(opt["dataValue"]))
    if opt.get("id"):
        return ("id", opt["id"])
    if opt.get("dataId"):
        return ("dataId", opt["dataId"])
    if opt.get("dataIndex") not in (None, ""):
        return ("dataIndex", str(opt["dataIndex"]))
    return ("text", (opt.get("text") or "").strip())


def _css_escape(s: str) -> str:
    # escape mínimo para IDs arbitrários
    return re.sub(r'(\\.#:[\\>+~*^$|])', r'\\\1', s or "")

def _best_key_for_option(opt: Dict[str, Any]) -> Tuple[str, Optional[str]]:
    """
    Decide qual atributo é mais estável para selecionar programaticamente.
    Preferência:
      1) value
      2) dataValue
      3) id
      4) dataId
      5) dataIndex
    Retorna (key_type, key_value) alinhado com nosso seletor.
    """
    if opt.get("value") not in (None, ""):
        return ("value", str(opt["value"]))
    if opt.get("dataValue") not in (None, ""):
        return ("dataValue", str(opt["dataValue"]))
    if opt.get("id"):
        return ("id", opt["id"])
    if opt.get("dataId"):
        return ("dataId", opt["dataId"])
    if opt.get("dataIndex") not in (None, ""):
        return ("dataIndex", str(opt["dataIndex"]))
    return ("text", (opt.get("text") or "").strip())

def _debug_dump_options(options, tag: str):
    try:
        Path(f"debug_opts_{tag}.json").write_text(
            json.dumps(options, ensure_ascii=False, indent=2), encoding="utf-8"
        )
    except Exception:
        pass


# ===== Localizador Universal v2 (IDs preferidos + label/atributos + fallback) =====
# IDs padrão do DOU para cada nível
LEVEL_IDS = {
    1: ["slcOrgs"],       # Organização Principal (Órgão)
    2: ["slcOrgsSubs"],   # Organização Subordinada / Unidade / Secretaria
    3: ["slcTipo"],       # Tipo do Ato
}

def _locate_root_by_id(frame, page, id_):
    """Tenta achar o root por ID no frame e, se não, na page."""
    # 1) frame
    try:
        loc = frame.locator(f"#{id_}")
        if loc.count() > 0 and loc.first.is_visible():
            return {"kind": "select", "sel": f"#{id_}", "index": 0, "handle": loc.first}
    except Exception:
        pass
    # 2) page (fora do frame)
    try:
        loc = page.locator(f"#{id_}")
        if loc.count() > 0 and loc.first.is_visible():
            return {"kind": "select", "sel": f"#{id_}", "index": 0, "handle": loc.first}
    except Exception:
        pass
    return None

def resolve_dropdown(level: int, frame, page, roots, label_regex: str = None, prefer_ids: list[str] | None = None):
    """
    Resolve o dropdown do 'level' (1,2,3) com máxima robustez:
      1) IDs preferidos (frame e page) -> retorna root;
      2) rótulo/atributos (label_regex) -> usa find_dropdown_by_label(...);
      3) fallback por índice do roots (0,1,2).
    """
    # 1) IDs preferidos
    id_list = (prefer_ids or []) + LEVEL_IDS.get(level, [])
    seen = set()
    for id_ in id_list:
        if not id_ or id_ in seen: 
            continue
        seen.add(id_)
        r = _locate_root_by_id(frame, page, id_)
        if r:
            return r

    # 2) rótulo/atributos
    if label_regex:
        cand = find_dropdown_by_label(frame, roots, label_regex)
        if cand:
            return cand

    # 3) fallback por índice
    idx = max(0, level - 1)
    if len(roots) > idx:
        return roots[idx]

    # fallback final: 1º disponível
    return roots[0] if roots else None


def sanitize_filename(name: str) -> str:
    name = re.sub(r'[\\/:*?"<>\|\r\n\t]+', "_", name)
    return name[:180].strip("_ ") or "out"

# ------------------------- URL helpers -------------------------
def origin_of(url: str) -> str:
    try:
        parsed = urlparse(url)
        return f"{parsed.scheme}://{parsed.netloc}"
    except Exception:
        return "https://www.in.gov.br"

def abs_url(base_or_page_url: str, href: str) -> str:
    if not href:
        return ""
    if href.startswith("http://") or href.startswith("https://"):
        return href
    base = origin_of(base_or_page_url)
    return urljoin(base + "/", href)

# ------------------------- Frames e roots -------------------------
def find_best_frame(context):
    page = context.pages[0]
    best = page.main_frame
    best_score = -1
    for fr in page.frames:
        score = 0
        try:
            for sel in DROPDOWN_ROOT_SELECTORS:
                score += fr.locator(sel).count()
        except Exception:
            pass
        if score > best_score:
            best_score = score
            best = fr
    return best

def collect_dropdown_roots(frame) -> List[Dict[str, Any]]:
    """
    Coleta raízes de dropdown (combobox/select/heurísticas), DEDUPLICA por id (quando houver),
    e prioriza select > combobox > unknown. Mantém ordenação por (y,x).
    """
    roots = []
    seen = set()

    def _maybe_add(sel: str, kind: str, loc):
        try:
            cnt = loc.count()
        except Exception:
            cnt = 0
        for i in range(cnt):
            h = loc.nth(i)
            try:
                box = h.bounding_box()
                if not box: 
                    continue
                key = (sel, i, round(box["y"],2), round(box["x"],2))
                if key in seen:
                    continue
                seen.add(key)
                roots.append({"selector": sel, "kind": kind, "index": i, "handle": h, "y": box["y"], "x": box["x"]})
            except Exception:
                continue

    # 1) ARIA combobox
    _maybe_add("role=combobox", "combobox", frame.get_by_role("combobox"))
    # 2) <select>
    _maybe_add("select", "select", frame.locator("select"))
    # 3) Heurísticas extras
    for sel in DROPDOWN_ROOT_SELECTORS:
        _maybe_add(sel, "unknown", frame.locator(sel))

    def _priority(kind:str)->int:
        return {"select": 3, "combobox": 2, "unknown": 1}.get(kind, 0)

    enriched = []
    for r in roots:
        h = r["handle"]
        try:
            el_id = h.get_attribute("id")
        except Exception:
            el_id = None
        enriched.append({**r, "id_attr": el_id})

    # DEDUPE por id; se sem id, usa (pos,selector)
    by_key = {}
    for r in enriched:
        if r.get("id_attr"):
            k = ("id", r["id_attr"])
        else:
            k = ("pos", round(r["y"],1), round(r["x"],1), r["selector"])
        best = by_key.get(k)
        if not best or _priority(r["kind"]) > _priority(best["kind"]):
            by_key[k] = r

    deduped = list(by_key.values())
    deduped.sort(key=lambda rr: (rr["y"], rr["x"]))
    return deduped

# ------------------------- Listbox / Combobox helpers -------------------------
def listbox_present(frame) -> bool:
    for sel in LISTBOX_SELECTORS:
        try:
            if frame.locator(sel).count() > 0:
                return True
        except Exception:
            pass
    page = frame.page
    for sel in LISTBOX_SELECTORS:
        try:
            if page.locator(sel).count() > 0:
                return True
        except Exception:
            pass
    return False

def get_listbox_container(frame):
    for sel in LISTBOX_SELECTORS:
        try:
            loc = frame.locator(sel)
            if loc.count() > 0:
                return loc.first
        except Exception:
            pass
    page = frame.page
    for sel in LISTBOX_SELECTORS:
        try:
            loc = page.locator(sel)
            if loc.count() > 0:
                return loc.first
        except Exception:
            pass
    return None

def open_dropdown(frame, root) -> bool:
    """
    Abre um dropdown visual (combobox custom ou <select> estilizado).
    Aceita: root dict {"handle": Locator} OU Locator direto.
    Estratégias: click normal -> force -> ícone interno -> teclado.
    Busca listbox no frame e também na page.
    """
    # Aceitar dict com "handle" ou Locator direto
    h = root["handle"] if isinstance(root, dict) else root

    def _listbox_present():
        if listbox_present(frame):
            return True
        # também procura na page (fora do frame)
        try:
            page = frame.page
            for sel in LISTBOX_SELECTORS:
                if page.locator(sel).count() > 0:
                    return True
        except Exception:
            pass
        return False

    # já aberto?
    if _listbox_present():
        return True

    # clique normal
    try:
        h.scroll_into_view_if_needed(timeout=2000)
        h.click(timeout=2500)
        frame.wait_for_timeout(120)
        if _listbox_present():
            return True
    except Exception:
        pass

    # clique forçado
    try:
        h.click(timeout=2500, force=True)
        frame.wait_for_timeout(120)
        if _listbox_present():
            return True
    except Exception:
        pass
    # seta/ícone interno
    try:
        arrow = h.locator("xpath=.//*[contains(@class,'arrow') or contains(@class,'icon') or contains(@class,'caret')]").first
        if arrow and arrow.count() > 0 and arrow.is_visible():
            arrow.click(timeout=2500)
            frame.wait_for_timeout(120)
            if _listbox_present():
                return True
    except Exception:
        pass

    # teclado
    try:
        h.focus()
        frame.page.keyboard.press("Enter"); frame.wait_for_timeout(120)
        if _listbox_present(): return True
        frame.page.keyboard.press("Space"); frame.wait_for_timeout(120)
        if _listbox_present(): return True
        frame.page.keyboard.press("Alt+ArrowDown"); frame.wait_for_timeout(120)
        if _listbox_present(): return True
    except Exception:
        pass
    return _listbox_present()


def scroll_listbox_all(container, frame) -> None:
    try:
        for _ in range(60):
            changed = container.evaluate(
                """el => { const b = el.scrollTop; el.scrollTop = el.scrollHeight; return el.scrollTop !== b; }"""
            )
            frame.wait_for_timeout(80)
            if not changed:
                break
    except Exception:
        for _ in range(15):
            try:
                frame.page.keyboard.press("End")
            except Exception:
                pass
            frame.wait_for_timeout(80)

def read_open_list_options(frame) -> List[Dict[str, Any]]:
    """
    Lê as opções do listbox aberto, varre provedores conhecidos,
    rola até o final, e CAPTURA também 'id' e 'data-id' quando existirem,
    além de text/value/dataValue/dataIndex.
    """
    container = get_listbox_container(frame)
    if not container:
        return []

    # rola para materializar itens virtualizados
    scroll_listbox_all(container, frame)

    options = []
    for sel in OPTION_SELECTORS:
        try:
            opts = container.locator(sel)
            k = opts.count()
        except Exception:
            k = 0

        for i in range(k):
            o = opts.nth(i)
            try:
                if not o.is_visible():
                    continue
                text = (o.text_content() or "").strip()
                val  = o.get_attribute("value")
                dv   = o.get_attribute("data-value")
                di   = o.get_attribute("data-index") or o.get_attribute("data-option-index") or str(i)
                oid  = o.get_attribute("id")
                did  = o.get_attribute("data-id") or o.get_attribute("data-key") or o.get_attribute("data-code")

                # guarda tudo que acharmos
                if text or val or dv or di or oid or did:
                    options.append({
                        "text": text,
                        "value": val,
                        "dataValue": dv,
                        "dataIndex": di,
                        "id": oid,
                        "dataId": did
                    })
            except Exception:
                pass

    # dedupe por (id, dataId, text, value, dataValue, dataIndex)
    seen = set()
    uniq = []
    for o in options:
        key = (o.get("id"), o.get("dataId"), o.get("text"), o.get("value"), o.get("dataValue"), o.get("dataIndex"))
        if key in seen:
            continue
        seen.add(key)
        uniq.append(o)

    try:
        frame.page.keyboard.press("Escape")
        frame.wait_for_timeout(100)
    except Exception:
        pass
    return uniq

# ------------------------- <select> helpers -------------------------
def is_select_root(root: Dict[str, Any]) -> bool:
    try:
        tag = root["handle"].evaluate("el => el.tagName && el.tagName.toLowerCase()")
        if tag == "select":
            return True
    except Exception:
        pass
    return root.get("selector") == "select"

def read_select_options(frame, root: Dict[str, Any]) -> List[Dict[str, Any]]:
    sel = root["handle"]
    try:
        return sel.evaluate(
            """
            el => Array.from(el.options).map((o,i)=>({
                text: (o.textContent||'').trim(),
                value: o.value,
                dataValue: o.getAttribute('data-value'),
                dataIndex: i
            }))
            """
        ) or []
    except Exception:
        return []

def select_by_key_select(frame, root: Dict[str, Any], key: str, key_type: str) -> bool:
    """
    Seleciona opção em <select> nativo com robustez máxima:
      - Rebusca o elemento por id (frame e page) – o DOM pode ser recriado após N1/N2.
      - Aguarda popular opções (até 8s).
      - Tenta select_option (label/value/index).
      - Fallback por texto normalizado (== ou contém).
      - Fallback "Todos" -> value='0'.
      - Fallback JS: selectedIndex/value + dispatch('input'/'change').
      - Fallback final: abre dropdown e clica em 'option' como se fosse combobox custom.
    """
    page = frame.page
    target = "" if key is None else str(key)
    nkey = normalize_text(target)

    # (Re)descobrir o <select> por id (root pode estar stale)
    sel = root.get("handle")
    try:
        _id = sel.get_attribute("id")
    except Exception:
        _id = None

    def _refetch_by_id(id_):
        if not id_:
            return None
        # frame
        try:
            loc = frame.locator(f"#{id_}")
            if loc.count() > 0 and loc.first.is_visible():
                return loc.first
        except Exception:
            pass
        # page
        try:
            loc = page.locator(f"#{id_}")
            if loc.count() > 0 and loc.first.is_visible():
                return loc.first
        except Exception:
            pass
        return None

    ref = _refetch_by_id(_id) if _id else None
    if ref is not None:
        sel = ref

    # Lê opções
    def _read_opts():
        try:
            return sel.evaluate("""
                el => Array.from(el.options || []).map((o,i) => ({
                    text: (o.label || o.textContent || '').trim(),
                    value: o.value,
                    dataValue: o.getAttribute('data-value'),
                    dataIndex: i
                }))
            """) or []
        except Exception:
            return []

    # Aguarda opções (até 8s)
    opts = _read_opts()
    waited = 0
    while waited < 8000 and len(opts) == 0:
        frame.wait_for_timeout(200)
        waited += 200
        opts = _read_opts()

    def _match_by_text(options, wanted_norm):
        for o in options:
            nt = normalize_text(o.get("text") or "")
            if nt == wanted_norm or (wanted_norm and wanted_norm in nt):
                return o
        return None

    # Tentativas
    try:
        if key_type == "value":
            sel.select_option(value=target)
            page.wait_for_load_state("networkidle", timeout=60_000)
            return True

        if key_type in ("dataValue", "dataIndex"):
            cand = None
            for o in opts:
                if key_type == "dataValue" and (o.get("dataValue") is not None) and str(o["dataValue"]) == target:
                    cand = o; break
                if key_type == "dataIndex" and str(o.get("dataIndex")) == target:
                    cand = o; break
            if cand:
                if cand.get("value") not in (None, ""):
                    sel.select_option(value=str(cand["value"]))
                else:
                    sel.select_option(index=int(cand["dataIndex"]))
                page.wait_for_load_state("networkidle", timeout=60_000)
                return True

        if key_type == "text":
            # 1) label "exata"
            try:
                sel.select_option(label=target)
                page.wait_for_load_state("networkidle", timeout=60_000)
                return True
            except Exception:
                pass
            # 2) match normalizado
            cand = _match_by_text(opts, nkey)
            if cand:
                if cand.get("value") not in (None, ""):
                    sel.select_option(value=str(cand["value"]))
                else:
                    sel.select_option(index=int(cand["dataIndex"]))
                page.wait_for_load_state("networkidle", timeout=60_000)
                return True
            # 3) "Todos" -> tentar value='0'
            if nkey == "todos":
                try:
                    sel.select_option(value="0")
                    page.wait_for_load_state("networkidle", timeout=60_000)
                    return True
                except Exception:
                    pass

        # 4) Fallback JS (selectedIndex + change)
        idx = -1
        try:
            idx = sel.evaluate(
                """
                (el, wnorm) => {
                  const opts = Array.from(el.options || []);
                  const norm = s => (s || '').normalize('NFKD')
                    .replace(/[^\x00-\x7F]/g,'')
                    .toLowerCase()
                    .replace(/\s+/g,' ')
                    .trim();
                  let found = -1;
                  for (let i=0; i<opts.length; i++){
                    const t = norm(opts[i].label || opts[i].textContent || '');
                    if (t === wnorm || (wnorm && t.includes(wnorm))) { found = i; break; }
                  }
                  if (found >= 0) {
                    el.selectedIndex = found;
                    el.dispatchEvent(new Event('input', {bubbles:true}));
                    el.dispatchEvent(new Event('change', {bubbles:true}));
                    return found;
                  }
                  return -1;
                }
                """,
                nkey
            )
        except Exception:
            idx = -1
        if idx is not None and int(idx) >= 0:
            page.wait_for_load_state("networkidle", timeout=60_000)
            return True

    except Exception:
        pass

    # 5) Fallback FINAL: tratar como combobox custom (abrir e clicar opção)
    try:
        # Abre o dropdown "visual"
        if open_dropdown(frame, sel):
            # procura por role=option com o texto
            try:
                opt = frame.get_by_role("option", name=re.compile(rf"^{re.escape(key)}$", re.I)).first
                if opt and opt.count() > 0 and opt.is_visible():
                    opt.click(timeout=4000)
                    page.wait_for_load_state("networkidle", timeout=60_000)
                    return True
            except Exception:
                pass
            # varrer containers conhecidos
            container = get_listbox_container(frame) or get_listbox_container(frame.page)
            if container:
                for selopt in OPTION_SELECTORS:
                    try:
                        opts_loc = container.locator(selopt)
                        k = opts_loc.count()
                    except Exception:
                        k = 0
                    for i in range(k):
                        o = opts_loc.nth(i)
                        try:
                            if not o.is_visible(): continue
                            txt = (o.text_content() or "").strip()
                            if normalize_text(txt) == nkey or (nkey and nkey in normalize_text(txt)):
                                o.click(timeout=4000)
                                page.wait_for_load_state("networkidle", timeout=60_000)
                                return True
                        except Exception:
                            pass
    except Exception:
        pass

    # dump de debug para investigação
    try:
        _debug_dump_options(opts, "n3_select_fail")
    except Exception:
        pass
    return False

# ------------------------- API unificada -------------------------
def label_for_control(frame, root: Dict[str, Any]) -> str:
    h = root["handle"]
    try:
        aria = h.get_attribute("aria-label")
        if aria:
            return aria.strip()
    except Exception:
        pass
    try:
        _id = h.get_attribute("id")
        if _id:
            lab = frame.evaluate(
                """
                (id) => {
                    const l = document.querySelector(`label[for="${id}"]`);
                    return l ? l.textContent.trim() : null;
                }
                """,
                _id,
            )
            if lab:
                return lab
    except Exception:
        pass
    try:
        prev = h.locator("xpath=preceding::label[1]").first
        if prev and prev.count() > 0 and prev.is_visible():
            t = (prev.text_content() or "").strip()
            if t:
                return t
    except Exception:
        pass
    return ""

def find_dropdown_by_label(frame, roots, label_regex: Optional[str]):
    """
    Localiza o dropdown pelo rótulo (<label>) OU por atributos do root
    (placeholder, id, name, aria-label). Faz match por regex e também
    por versão normalizada (sem acento, minúsculas).
    Se nada bater, aplica fallback específico do DOU por id conhecido.
    """
    if not label_regex:
        return None

    pat = re.compile(label_regex, re.I)
    want_norm = normalize_text(label_regex)

    # 1) label formal (<label>)
    for r in roots:
        lab = label_for_control(frame, r) or ""
        if pat.search(lab):
            return r

    # 2) atributos do root
    def _attrs_text(h):
        vals = []
        for a in ("placeholder", "id", "name", "aria-label"):
            try:
                v = h.get_attribute(a) or ""
            except Exception:
                v = ""
            if v:
                vals.append(v)
        return " | ".join(vals)

    for r in roots:
        h = r["handle"]
        txt = _attrs_text(h)
        if txt and pat.search(txt):
            return r

    # 3) normalizado (sem acento/minúsculo)
    if want_norm:
        for r in roots:
            h = r["handle"]
            txt = normalize_text(_attrs_text(h))
            if want_norm in txt:
                return r

    # 4) Fallback específico DOU por id conhecido
    pref = normalize_text(label_regex)
    prefer_tipo = ("tipo" in pref)
    prefer_sub  = ("sub" in pref) or ("unidad" in pref) or ("secretar" in pref)
    order = []
    if prefer_tipo:
        order = ["slcTipo", "slcOrgsSubs", "slcOrgs"]
    elif prefer_sub:
        order = ["slcOrgsSubs", "slcOrgs", "slcTipo"]
    else:
        order = ["slcOrgs", "slcOrgsSubs", "slcTipo"]

    for wanted in order:
        for r in roots:
            try:
                if (r["handle"].get_attribute("id") or "") == wanted:
                    return r
            except Exception:
                pass

    return None

def read_dropdown_options(frame, root: Dict[str, Any]) -> List[Dict[str, Any]]:
    if is_select_root(root):
        return read_select_options(frame, root)
    if not open_dropdown(frame, root):
        return []
    return read_open_list_options(frame)

def select_by_key(frame, root: Dict[str, Any], key: str, key_type: str) -> bool:
    """
    Seleciona por chave em qualquer tipo de dropdown:

    - <select>: usa select_by_key_select() (label/value/index) + fallbacks.
    - Combobox custom: abre listbox, coleta opções (com id/data-id/dv/di),
      1) tenta bater por key_type (quando fizer sentido),
      2) se não der, casa por TEXTO normalizado,
      3) efetiva a seleção por ID/atributo estável (id/data-id/value/data-value/data-index),
         clicando a opção ou setando via evento,
      4) fallback final por clique no texto.
    """
    page = frame.page

    # 1) <select> nativo?
    if is_select_root(root):
        if select_by_key_select(frame, root, key, key_type):
            return True
        # fallback: tratar como custom
        try:
            h = root["handle"]
            if open_dropdown(frame, h):
                container = get_listbox_container(frame) or get_listbox_container(page)
                if container:
                    # usa o mesmo fluxo do custom logo abaixo
                    # coletar opções enriquecidas
                    opts = read_open_list_options(frame)
                    return _select_in_custom_from_options(frame, container, opts, key, key_type)
        except Exception:
            pass
        try: page.keyboard.press("Escape")
        except Exception: pass
        return False

    # 2) Combobox custom
    if not open_dropdown(frame, root["handle"]):
        return False
    container = get_listbox_container(frame) or get_listbox_container(page)
    if not container:
        return False

    # coletar opções enriquecidas (com id/data-id/value/dv/di)
    opts = read_open_list_options(frame)
    return _select_in_custom_from_options(frame, container, opts, key, key_type)

def _select_in_custom_from_options(frame, container, options: List[Dict[str, Any]], key: str, key_type: str) -> bool:
    """
    Dado um conjunto de opções (com id/data-id/value/dataValue/dataIndex/text),
    determina a melhor candidata e seleciona priorizando identificadores estáveis.
    """
    page = frame.page
    nkey = normalize_text(str(key))

    # 1) tentar bater por key_type fornecido (value/dataValue/dataIndex/text)
    def _match_by_key_type(o):
        if key_type == "value" and o.get("value") not in (None, ""):
            return str(o["value"]) == str(key)
        if key_type == "dataValue" and o.get("dataValue") not in (None, ""):
            return str(o["dataValue"]) == str(key)
        if key_type == "dataIndex" and o.get("dataIndex") not in (None, ""):
            return str(o["dataIndex"]) == str(key)
        if key_type == "text":
            return normalize_text(o.get("text") or "") == nkey
        return False

    picks = [o for o in options if _match_by_key_type(o)]

    # 2) se nada bateu, usa TEXTO normalizado (mais robusto para custom)
    if not picks:
        for o in options:
            if nkey and nkey == normalize_text(o.get("text") or ""):
                picks = [o]; break
        # como fallback, 'contains'
        if not picks:
            for o in options:
                if nkey and nkey in normalize_text(o.get("text") or ""):
                    picks = [o]; break

    if not picks:
        return False

    target = picks[0]
    ktype, kval = _best_key_for_option(target)

    # 3) tentar selecionar por ID/atributo estável
    try:
        if ktype == "id" and kval:
            opt = container.locator(f"##{_css_escape(kval)}")
            if opt and opt.count() > 0 and opt.first.is_visible():
                opt.first.click(timeout=4000)
                page.wait_for_load_state("networkidle", timeout=60_000)
                return True

        if ktype == "dataId" and kval:
            opt = container.locator(f"[data-id='{_css_escape(kval)}'], [data-key='{_css_escape(kval)}'], [data-code='{_css_escape(kval)}']")
            if opt and opt.count() > 0 and opt.first.is_visible():
                opt.first.click(timeout=4000)
                page.wait_for_load_state("networkidle", timeout=60_000)
                return True

        if ktype == "value" and kval not in (None, ""):
            # alguns custom mantém 'value' em atributo do item
            opt = container.locator(f"[value='{_css_escape(kval)}']")
            if opt and opt.count() > 0 and opt.first.is_visible():
                opt.first.click(timeout=4000)
                page.wait_for_load_state("networkidle", timeout=60_000)
                return True

        if ktype == "dataValue" and kval not in (None, ""):
            opt = container.locator(f"[data-value='{_css_escape(kval)}']")
            if opt and opt.count() > 0 and opt.first.is_visible():
                opt.first.click(timeout=4000)
                page.wait_for_load_state("networkidle", timeout=60_000)
                return True

        if ktype == "dataIndex" and kval not in (None, ""):
            # alguns frameworks usam data-index estável
            opt = container.locator(f"[data-index='{_css_escape(kval)}'], [data-option-index='{_css_escape(kval)}']")
            if opt and opt.count() > 0 and opt.first.is_visible():
                opt.first.click(timeout=4000)
                page.wait_for_load_state("networkidle", timeout=60_000)
                return True
    except Exception:
        pass

    # 4) fallback final: clicar pelo texto visível
    try:
        # match exato
        opt = container.get_by_role("option", name=re.compile(rf"^{re.escape(str(target.get('text') or ''))}$", re.I)).first
        if opt and opt.count() > 0 and opt.is_visible():
            opt.click(timeout=4000)
            page.wait_for_load_state("networkidle", timeout=60_000)
            return True
    except Exception:
        pass

    try:
        # match contains (normalizado)
        nk = normalize_text(target.get("text") or "")
        any_opt = container.locator(
            "xpath=//*[contains(translate(normalize-space(text()), 'ABCDEFGHIJKLMNOPQRSTUVWXYZ','abcdefghijklmnopqrstuvwxyz'), $k)]",
            k=nk
        ).first
        if any_opt and any_opt.count() > 0 and any_opt.is_visible():
            any_opt.click(timeout=4000)
            page.wait_for_load_state("networkidle", timeout=60_000)
            return True
    except Exception:
        pass
    try:
        page.keyboard.press("Escape")
    except Exception:
        pass
    return False

# ------------------------- Busca (query) -------------------------
def apply_query(frame, query: Optional[str]) -> None:
    if not query:
        return
    locs = [
        lambda f: f.locator("#search-bar").first,
        lambda f: f.get_by_role("searchbox").first,
        lambda f: f.get_by_role("textbox", name=re.compile("pesquis", re.I)).first,
        lambda f: f.locator('input[placeholder*="esquis" i]').first,
        lambda f: f.locator('input[type="search"]').first,
        lambda f: f.locator('input[aria-label*="pesquis" i]').first,
        lambda f: f.locator('input[id*="pesquis" i], input[name*="pesquis" i]').first,
    ]
    sb = None
    for get in locs:
        try:
            cand = get(frame)
            if cand and cand.count() > 0:
                cand.wait_for(state="visible", timeout=5000)
                sb = cand
                break
        except Exception:
            continue
    if not sb:
        print("[Aviso] Campo de busca não encontrado; seguindo sem query.")
        return
    try:
        sb.scroll_into_view_if_needed(timeout=1500)
    except Exception:
        pass
    try:
        sb.fill(query)
    except PWTimeout:
        sb.evaluate(
            "(el, val) => { el.value = val; el.dispatchEvent(new Event('input',{bubbles:true})); }",
            query
        )
    try:
        sb.press("Enter")
    except Exception:
        try:
            bt = frame.get_by_role("button", name=re.compile("pesquis", re.I))
            if bt.count() > 0 and bt.first.is_visible():
                bt.first.click()
        except Exception:
            pass
    frame.page.wait_for_load_state("networkidle", timeout=90_000)

# ------------------------- Coleta de links -------------------------
def collect_links(
    frame,
    max_links: int,
    max_scrolls: int = 40,
    scroll_pause_ms: int = 350,
    stable_rounds: int = 3,
) -> List[Dict[str, str]]:
    page = frame.page
    container = page.locator("#hierarchy_content")
    anchors = container.locator('a[href*="/web/dou/"]') if container.count() > 0 else page.locator('a[href*="/web/dou/"]')
    last = -1
    stable = 0
    for _ in range(max_scrolls):
        try:
            count = anchors.count()
        except Exception:
            count = 0
        if count >= max_links:
            break
        if count == last:
            stable += 1
        else:
            stable = 0
        if stable >= stable_rounds:
            break
        last = count
        try:
            page.evaluate("window.scrollTo(0, document.body.scrollHeight)")
        except Exception:
            pass
        page.wait_for_timeout(scroll_pause_ms)

    try:
        items = page.evaluate(f"""
            () => {{
                const root = document.querySelector('#hierarchy_content') || document;
                const anchors = Array.from(root.querySelectorAll('a[href*="/web/dou/"]'));
                return anchors.slice(0, {max_links}).map(a => ({{
                    titulo: (a.textContent || '').trim(),
                    link: a.getAttribute('href') || ''
                }})).filter(x => x.titulo && x.link);
            }}
        """)
    except Exception:
        items = []
    return items

# ------------------------- Detalhamento -------------------------
def text_of(locator) -> str:
    try:
        if locator and locator.count() > 0 and locator.first.is_visible():
            t = locator.first.text_content() or ""
            return re.sub(r"\s+", " ", t).strip()
    except Exception:
        pass
    return ""

def meta_content(page, selector: str) -> Optional[str]:
    try:
        el = page.locator(selector).first
        if el and el.count() > 0:
            v = el.get_attribute("content")
            if v:
                return v.strip()
    except Exception:
        pass
    return None

def find_dt_dd_value(page, labels_regex: str) -> Optional[str]:
    pat = re.compile(labels_regex, re.I)
    dts = page.locator("dl dt")
    try:
        n = dts.count()
    except Exception:
        n = 0
    for i in range(min(n, 400)):
        try:
            dt = dts.nth(i)
            if not dt.is_visible():
                continue
            txt = (dt.text_content() or "").strip()
            if not txt or not pat.search(txt):
                continue
            dd = dt.locator("xpath=following-sibling::dd[1]")
            val = text_of(dd)
            if val:
                return val
        except Exception:
            continue
    cands = page.locator("strong, b, label, span")
    try:
        k = cands.count()
    except Exception:
        k = 0
    for i in range(min(k, 600)):
        try:
            c = cands.nth(i)
            if not c.is_visible():
                continue
            t = (c.text_content() or "").strip()
            if not t or not pat.search(t):
                continue
            parent = c.locator("xpath=..")
            val = text_of(parent)
            if val:
                val2 = re.sub(r"^\s*(Órgão|Orgao|Tipo|Tipo do Ato)\s*:\s*", "", val, flags=re.I).strip()
                if val2:
                    return val2
        except Exception:
            continue
    return None

def collect_article_text(page, max_chars: int = 6000) -> str:
    """Coleta texto do artigo unificado, sem quebrar demais, para base de resumo."""
    selectors = [
        "article", "main article", "div[class*=materia]", "main"
    ]
    buf = []
    for sel in selectors:
        loc = page.locator(sel)
        try:
            if loc.count() > 0 and loc.first.is_visible():
                # pega parágrafos
                ps = loc.locator("p")
                n = min(ps.count(), 300)
                for i in range(n):
                    t = (ps.nth(i).text_content() or "").strip()
                    if t:
                        buf.append(re.sub(r"\s+", " ", t))
                if buf:
                    break
        except Exception:
            continue
    text = " ".join(buf)
    text = re.sub(r"\s+", " ", text).strip()
    if len(text) > max_chars:
        text = text[:max_chars].rsplit(".", 1)[0] + "."
    return text


def scrape_detail(context, url: str, timeout_ms: int = 60_000) -> Dict[str, Any]:
    page = context.new_page()
    data = {
        "detail_url": url,
        "titulo": None,
        "ementa": None,
        "orgao": None,
        "tipo_ato": None,
        "secao": None,
        "data_publicacao": None,
        "pdf_url": None,
        "edicao": None,
        "pagina": None,
    }
    try:
        page.set_default_timeout(timeout_ms)
        page.goto(url, wait_until="domcontentloaded")
        page.wait_for_load_state("networkidle", timeout=timeout_ms)
        close_cookies(page)

        # Título
        title_candidates = [
            lambda: meta_content(page, 'meta[property="og:title"]'),
            lambda: meta_content(page, 'meta[name="dc.title"]'),
        ]
        titulo = None
        for get in title_candidates:
            titulo = get()
            if titulo:
                break
        if not titulo:
            for sel in [
                "article h1", "main article h1", "h1",
                "article h2", "main article h2", "h2",
                "[class*=titulo] h1, [class*=title] h1", "[class*=titulo], [class*=title]"
            ]:
                t = text_of(page.locator(sel))
                if t:
                    titulo = t
                    break
        if not titulo:
            titulo = (page.title() or "").strip() or None
        data["titulo"] = titulo

        # Data de publicação
        date_candidates = [
            lambda: meta_content(page, 'meta[property="article:published_time"]'),
            lambda: meta_content(page, 'meta[name="publicationDate"]'),
            lambda: meta_content(page, 'meta[name="dc.date"]'),
        ]
        pub = None
        for get in date_candidates:
            pub = get()
            if pub:
                break
        if not pub:
            time_el = page.locator("time[datetime]").first
            if time_el and time_el.count() > 0:
                try:
                    pub = time_el.get_attribute("datetime")
                except Exception:
                    pass
        def norm_date(s: Optional[str]) -> Optional[str]:
            if not s:
                return None
            ss = s.strip()
            m = re.match(r"(\d{4})-(\d{2})-(\d{2})", ss)
            if m:
                return f"{m.group(1)}-{m.group(2)}-{m.group(3)}"
            m = re.match(r"(\d{2})/(\d{2})/(\d{4})", ss)
            if m:
                return f"{m.group(3)}-{m.group(2)}-{m.group(1)}"
            return ss

        data["data_publicacao"] = norm_date(pub)

        # Órgão / Tipo do ato
        if not data["orgao"]:
            data["orgao"] = find_dt_dd_value(page, r"(Órgão|Orgao)")
        if not data["tipo_ato"]:
            data["tipo_ato"] = find_dt_dd_value(page, r"(Tipo|Tipo do Ato)")

        # Seção
        if not data["secao"]:
            try:
                sec_link = page.locator('a[href*="secao=DO"]').first
                stext = text_of(sec_link)
                if stext:
                    data["secao"] = stext
            except Exception:
                pass
        if not data["secao"]:
            alt = meta_content(page, 'meta[name="dc.subject"]')
            if alt and re.search(r"DO[123]", alt):
                data["secao"] = alt

        # Ementa
        for sel in ["article .texto p", "article p", "main article p", "div[class*=materia] p", "main p"]:
            p = text_of(page.locator(sel))
            if p and len(p) > 0:
                data["ementa"] = p
            # Texto completo (para resumo)
            try:
                full_txt = collect_article_text(page, max_chars=8000)
                if full_txt and (not data.get("ementa") or len(full_txt) > len(data["ementa"])):
                    data["texto"] = full_txt
                else:
                    data["texto"] = data.get("ementa") or ""
            except Exception:
                data["texto"] = data.get("ementa") or ""
                break
        # PDF
        pdf = None
        try:
            pdfs = page.locator("a[href$='.pdf'], a[href*='.pdf?']")
            kk = pdfs.count()
        except Exception:
            kk = 0
        for i in range(min(kk, 30)):
            a = pdfs.nth(i)
            try:
                href = a.get_attribute("href")
                if href and (href.lower().endswith(".pdf") or ".pdf?" in href.lower()):
                    pdf = abs_url(page.url, href)
                    break
            except Exception:
                pass
        if not pdf:
            try:
                vc = page.get_by_role("link", name=re.compile("VERS(Ã|A)O CERTIFICADA", re.I)).first
                if vc and vc.count() > 0 and vc.is_visible():
                    href = vc.get_attribute("href")
                    if href:
                        pdf = abs_url(page.url, href)
            except Exception:
                pass
        data["pdf_url"] = pdf

        # Edicao/Pagina
        try:
            body_text = text_of(page.locator("article")) or text_of(page.locator("main")) or (page.inner_text("body") or "")
            m_ed = re.search(r"Edi[cç][aã]o\s*:\s*(\d+)", body_text, re.I)
            m_pg = re.search(r"P[aá]gina\s*:\s*(\d+)", body_text, re.I)
            if m_ed:
                data["edicao"] = m_ed.group(1)
            if m_pg:
                data["pagina"] = m_pg.group(1)
        except Exception:
            pass
    except Exception:
        pass
    finally:
        try:
            page.close()
        except Exception:
            pass

    base = (url or "") + "\n\n" + (data.get("titulo") or "")
    data["hash"] = hashlib.sha1(base.encode("utf-8")).hexdigest()
    return data

# ------------------------- Debug -------------------------
def dump_debug(page, prefix="debug") -> None:
    try:
        page.screenshot(path=f"{prefix}.png", full_page=True)
    except Exception:
        pass
    try:
        html = page.content()
        Path(f"{prefix}.html").write_text(html, encoding="utf-8")
    except Exception:
        pass
# ------------------------- Impressão (list) -------------------------
def print_list(label: str, level: int, options: List[Dict[str, Any]]) -> None:
    print(f"\n[Dropdown {level}] {label or '(sem rótulo)'} — {len(options)} opções")
    print("-" * 110)
    for i, o in enumerate(options, 1):
        print(
            f"{i:>2} "
            f"text='{o.get('text','')}' "
            f"value={o.get('value')} "
            f"data-value={o.get('dataValue')} "
            f"data-index={o.get('dataIndex')}"
        )
    print("-" * 110)

# ------------------------- Fluxos (list / run) -------------------------
def flow_list(
    context, data, secao, level,
    key1, key1_type, key2, key2_type,
    out_path, debug_dump, label1=None, label2=None, label3=None
):
    page = context.pages[0]
    goto(page, data, secao)
    frame = find_best_frame(context)
    roots = collect_dropdown_roots(frame)
    if not roots:
        print("[Erro] Nenhum dropdown detectado.")
        if debug_dump:
            dump_debug(page, "debug_list_no_roots")
        sys.exit(1)

    # --- Nível 1 ---
    r1 = resolve_dropdown(1, frame, page, roots, label1)
    lab1 = label_for_control(frame, r1)
    opts1 = read_dropdown_options(frame, r1)
    print_list(lab1, 1, opts1)

    if level >= 2:
        if not key1 or not key1_type:
            print("[Erro] Para listar nível 2, informe --key1-type e --key1.")
            sys.exit(1)
        if not select_by_key(frame, r1, key1, key1_type):
            print(f"[Erro] Nível 1: não consegui selecionar {key1_type}='{key1}'.")
            if debug_dump:
                dump_debug(page, "debug_list_select1_fail")
            sys.exit(1)
        page.wait_for_load_state("networkidle", timeout=90_000)
        roots = collect_dropdown_roots(frame)
        if len(roots) < 1:
            print("[Erro] Nível 2 não apareceu após selecionar o nível 1.")
            if debug_dump:
                dump_debug(page, "debug_list_no_level2")
            sys.exit(1)
        r2 = resolve_dropdown(2, frame, page, roots, label2)
        lab2 = label_for_control(frame, r2)
        opts2 = read_dropdown_options(frame, r2)
        print_list(lab2, 2, opts2)

    if level >= 3:
        if not key2 or not key2_type:
            print("[Erro] Para listar nível 3, informe --key2-type e --key2.")
            sys.exit(1)
        if not select_by_key(frame, r2, key2, key2_type):
            print(f"[Erro] Nível 2: não consegui selecionar {key2_type}='{key2}'.")
            if debug_dump:
                dump_debug(page, "debug_list_select2_fail")
            sys.exit(1)
        page.wait_for_load_state("networkidle", timeout=90_000)
        roots = collect_dropdown_roots(frame)
        if len(roots) < 1:
            print("[Erro] Nível 3 não apareceu após selecionar o nível 2.")
            if debug_dump:
                dump_debug(page, "debug_list_no_level3")
            sys.exit(1)
        r3 = resolve_dropdown(3, frame, page, roots, label3)
        lab3 = label_for_control(frame, r3)
        opts3 = read_dropdown_options(frame, r3)
        print_list(lab3, 3, opts3)

    payload = {"data": data, "secao": secao}
    if level == 1:
        payload.update({"level": 1, "label": lab1, "options": opts1})
    elif level == 2:
        payload.update({"level": 2, "label1": lab1, "label2": lab2,
                        "key1": key1, "key1_type": key1_type, "options": opts2})
    else:
        payload.update({"level": 3, "label1": lab1, "label2": lab2, "label3": lab3,
                        "key1": key1, "key1_type": key1_type, "key2": key2, "key2_type": key2_type,
                        "options": opts3})
    Path(out_path).write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"\n[OK] Opções do nível {level} salvas em: {out_path}")


def flow_run(
    context, data, secao,
    key1, key1_type, key2, key2_type, key3, key3_type,
    query, max_links, out_path, debug_dump,
    scrape_details: bool, detail_timeout: int, fallback_date_if_missing: bool,
    label1: Optional[str] = None, label2: Optional[str] = None, label3: Optional[str] = None,
    max_scrolls: int = 40, scroll_pause_ms: int = 350, stable_rounds: int = 3,
    state_file: Optional[str] = None, bulletin: Optional[str] = None, bulletin_out: Optional[str] = None
):
    """Execução completa; resolve L1/L2/L3 com o mesmo algoritmo; L3 opcional; binds por ID no frame e na page."""
    page = context.pages[-1] if context.pages else context.new_page()
    result = {
        "data": data,
        "secao": secao,
        "selecoes": [
            {"level": 1, "type": key1_type, "key": key1},
            {"level": 2, "type": key2_type, "key": key2},
            {"level": 3, "type": key3_type, "key": key3},
        ],
        "query": query,
        "total": 0,
        "itens": [],
        "enriquecido": False,
    }

    def _save_and_return():
        # garantir pasta do JSON
        Path(out_path).parent.mkdir(parents=True, exist_ok=True)
        Path(out_path).write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
        print(f"[OK] Links salvos em: {out_path} (total={result['total']})")

    try:
        goto(page, data, secao)
        frame = find_best_frame(context)
        roots = collect_dropdown_roots(frame)
        if len(roots) < 1:
            print(f"[Info] Nenhuma edição/roots detectados para {data} {secao}. Pulando.")
            _save_and_return(); return

        # --- Nível 1 ---
        roots = collect_dropdown_roots(frame)
        r1 = resolve_dropdown(1, frame, page, roots, label1)
        if not r1 or not select_by_key(frame, r1, key1, key1_type):
            print(f"[Info] Nível 1: não consegui selecionar {key1_type}='{key1}'. Pulando.")
            _save_and_return(); return
        page.wait_for_load_state("networkidle", timeout=90_000)

        # --- Nível 2 ---
        roots = collect_dropdown_roots(frame)
        if len(roots) < 1:
            print("[Info] Nível 2: roots ausentes após N1. Pulando.")
            _save_and_return(); return
        r2 = resolve_dropdown(2, frame, page, roots, label2)
        if not r2 or not select_by_key(frame, r2, key2, key2_type):
            print(f"[Info] Nível 2: não consegui selecionar {key2_type}='{key2}'. Pulando.")
            _save_and_return(); return
        page.wait_for_load_state("networkidle", timeout=90_000)

        # --- Nível 3 (OPCIONAL) ---
        if key3 and key3_type:
            roots = collect_dropdown_roots(frame)
            r3 = resolve_dropdown(3, frame, page, roots, label3)
            if r3:
                if not select_by_key(frame, r3, key3, key3_type):
                    print(f"[Info] Nível 3: não consegui selecionar {key3_type}='{key3}'. Pulando N3 e seguindo.")
                else:
                    page.wait_for_load_state("networkidle", timeout=90_000)
            else:
                print("[Info] Nível 3: root não encontrado. Pulando N3 e seguindo.")
        else:
            print("[Info] Nível 3 não informado — executando com L1+L2 apenas.")

        # --- Query + coleta ---
        apply_query(frame, query)
        itens = collect_links(frame, max_links, max_scrolls=max_scrolls, scroll_pause_ms=scroll_pause_ms, stable_rounds=stable_rounds)

        # --- Enriquecimento ---
        enriched = itens
        if scrape_details:
            print(f"[Info] Enriquecendo {len(itens)} itens (timeout={detail_timeout} ms cada)...")
            enriched = []
            for it in itens:
                detail_url = abs_url(page.url, it.get("link"))
                try:
                    meta = scrape_detail(context, detail_url, timeout_ms=detail_timeout)
                except Exception:
                    meta = {"detail_url": detail_url}

                def is_bad_site_title(t: str) -> bool:
                    return normalize_text(t) in {"imprensa nacional", ""}

                titulo_list = it.get("titulo") or ""
                titulo_det = (meta.get("titulo") or "").strip()
                final_title = titulo_det if not is_bad_site_title(titulo_det) else titulo_list

                data_pub = meta.get("data_publicacao")
                if not data_pub and fallback_date_if_missing:
                    try:
                        dt = datetime.strptime(data, "%d-%m-%Y").date()
                        data_pub = dt.strftime("%Y-%m-%d")
                    except Exception:
                        pass

                enriched.append({
                    **it, **meta,
                    "titulo_listagem": titulo_list,
                    "titulo_detalhe": titulo_det,
                    "titulo": final_title,
                    "detail_url": detail_url,
                    "data_publicacao": data_pub or meta.get("data_publicacao"),
                    "data_publicacao_fallback": (not meta.get("data_publicacao")) if scrape_details else False,
                })

        # --- Dedup ---
        filtered = enriched
        if state_file:
            seen = set()
            state_path = Path(state_file)
            try:
                if state_path.exists():
                    for line in state_path.read_text(encoding="utf-8").splitlines():
                        try:
                            obj = json.loads(line)
                            if "hash" in obj: seen.add(obj["hash"])
                        except Exception: pass
            except Exception: pass

            new_items = []
            for it in enriched:
                h = it.get("hash")
                if h and h in seen: continue
                new_items.append(it)
                if h: seen.add(h)
            filtered = new_items

            try:
                if filtered:
                    with state_path.open("a", encoding="utf-8") as f:
                        for it in filtered:
                            if it.get("hash"):
                                f.write(json.dumps({"hash": it["hash"]}, ensure_ascii=False) + "\n")
            except Exception:
                pass

        # --- Finaliza ---
        result["total"] = len(filtered)
        result["itens"] = filtered
        result["enriquecido"] = bool(scrape_details)

        _save_and_return()

        if bulletin and bulletin_out:
            # garantir pasta do boletim
            Path(bulletin_out).parent.mkdir(parents=True, exist_ok=True)
            try:
                gen_bulletin(result, bulletin, bulletin_out)
                print(f"[OK] Boletim gerado: {bulletin_out}")
            except Exception as e:
                print(f"[Aviso] Falha ao gerar boletim: {e}")

    except Exception as e:
        print(f"[Aviso] Falha inesperada no run: {e}")
        if debug_dump:
            try: dump_debug(page, "debug_run_unexpected")
            except Exception: pass
        _save_and_return()

# ------------------------- Boletim -------------------------
def gen_bulletin(result: Dict[str, Any], kind: str, out_path: str) -> None:
    """
    Gera boletim agrupado por órgão/tipo.
    Suporta: "md", "html" e "docx".
    """
    from collections import defaultdict
    from pathlib import Path

    # garantir pasta do boletim
    Path(out_path).parent.mkdir(parents=True, exist_ok=True)

    groups = defaultdict(list)
    for it in result.get("itens", []):
        key = (it.get("orgao") or "Sem órgão", it.get("tipo_ato") or "Sem tipo")
        groups[key].append(it)

    date = result.get("data") or ""
    secao = result.get("secao") or ""

    def _mk_suffix(it: Dict[str, Any]) -> str:
        extra = []
        if it.get("data_publicacao"): extra.append(it["data_publicacao"])
        if it.get("secao"): extra.append(it["secao"])
        if it.get("edicao"): extra.append(f"Edição {it['edicao']}")
        if it.get("pagina"): extra.append(f"p. {it['pagina']}")
        return (" — " + " • ".join(extra)) if extra else ""

    # ---- DOCX ----
    if kind == "docx":
        try:
            from docx import Document
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            from docx.opc.constants import RELATIONSHIP_TYPE as RT
        except Exception as e:
            raise RuntimeError("Saída DOCX requer o pacote 'python-docx'. Instale com: pip install python-docx") from e

        def _add_hyperlink(paragraph, url: str, text: str, color="0000FF", underline=True):
            r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)
            new_run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            if color:
                c = OxmlElement('w:color'); c.set(qn('w:val'), color); rPr.append(c)
            if underline:
                u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
            new_run.append(rPr)
            t = OxmlElement('w:t'); t.text = text; new_run.append(t)
            hyperlink.append(new_run)
            paragraph._p.append(hyperlink)

        doc = Document()
        doc.add_heading(f"Boletim DOU — {date} ({secao})", 0)
        for (org, tipo), arr in groups.items():
            doc.add_heading(f"{org} — {tipo}", level=1)
            for it in arr:
                titulo = it.get("titulo") or it.get("titulo_listagem") or "Sem título"
                durl = it.get("detail_url") or it.get("link") or ""
                pdf = it.get("pdf_url") or ""
                suffix = _mk_suffix(it)

                p = doc.add_paragraph(style="List Bullet")
                if durl: _add_hyperlink(p, durl, titulo)
                else: p.add_run(titulo)
                if pdf:
                    p.add_run(" ["); _add_hyperlink(p, pdf, "PDF"); p.add_run("]")
                if suffix: p.add_run(suffix)

                base_text = it.get("texto") or it.get("ementa") or ""
                if base_text:
                    try:
                        snippet = summarize_text(base_text, max_lines=SUMMARY_LINES, keywords=SUMMARY_KEYWORDS, mode=SUMMARY_MODE)
                    except Exception:
                        snippet = base_text
                    snippet = (snippet or "").strip()
                    if snippet:
                        pr = doc.add_paragraph()
                        r = pr.add_run("Resumo: "); r.bold = True
                        pr.add_run(snippet)
        doc.save(out_path)
        return
    
    # ---- MARKDOWN ----
    if kind == "md":
        lines = [f"# Boletim DOU — {date} ({secao})", ""]
        for (org, tipo), arr in groups.items():
            lines += [f"## {org} — {tipo}", ""]
            for it in arr:
                t = it.get("titulo") or it.get("titulo_listagem") or ""
                durl = it.get("detail_url") or it.get("link") or ""
                pdf = it.get("pdf_url") or ""
                suffix = _mk_suffix(it)
                pdfpart = f" · PDF" if pdf else ""
                link = f"{t}" if durl else t
                lines.append(f"- {link}{pdfpart}{suffix}")
                base_text = it.get("texto") or it.get("ementa") or ""
                if base_text:
                    try:
                        snippet = summarize_text(base_text, max_sentences=SUMMARY_SENTENCES, keywords=SUMMARY_KEYWORDS)
                    except Exception:
                        snippet = base_text
                    snippet = (snippet or "").strip()
                    if snippet:
                        lines.append(f"  \n _Resumo:_ {snippet}")
                lines.append("")
        Path(out_path).write_text("\n".join(lines), encoding="utf-8")
        return
    # ---- HTML ----
    import html as htmllib
    parts = [f"<h1>Boletim DOU — {htmllib.escape(date)} ({htmllib.escape(secao)})</h1>"]
    for (org, tipo), arr in groups.items():
        parts.append(f"<h2>{htmllib.escape(org)} — {htmllib.escape(tipo)}</h2><ul>")
        for it in arr:
            t = it.get("titulo") or it.get("titulo_listagem") or ""
            durl = it.get("detail_url") or it.get("link") or ""
            pdf = it.get("pdf_url") or ""
            suffix = _mk_suffix(it)
            title = htmllib.escape(t)
            link = f'{htmllib.escape(durl)}{title}</a>' if durl else title
            pdfhtml = f' · {htmllib.escape(pdf)}PDF</a>' if pdf else ""
            parts.append(f"<li>{link}{pdfhtml}{htmllib.escape(suffix)}</li>")
        parts.append("</ul>")
    Path(out_path).write_text("\n".join(parts), encoding="utf-8")
    return
    
# ------------------------- PLANEJAMENTO A PARTIR DO MAPA (00) -------------------------
def _parse_secao_data_from_url(url: str) -> Tuple[Optional[str], Optional[str]]:
    secao = None; data = None
    try:
        if "?" in url:
            _, qs = url.split("?", 1)
            parts = [p for p in qs.split("&") if "=" in p]
            qd = dict(p.split("=", 1) for p in parts)
            secao = qd.get("secao")
            data = qd.get("data")
            if data:
                try:
                    _ = datetime.strptime(data, "%d-%m-%Y")
                except Exception:
                    data = None
    except Exception:
        pass
    return secao, data

def _choose_root_by_label(dropdowns: List[Dict[str, Any]], label_regex: Optional[str], fallback_index: int) -> Dict[str, Any]:
    """
    Escolhe o root no JSON do mapa (modo PLAN), preferindo:
      1) label que case com label_regex;
      2) id típico do DOU (slcOrgs / slcOrgsSubs / slcTipo), inferido do label_regex;
      3) fallback por índice.
    """
    if not dropdowns:
        raise RuntimeError("Mapa sem dropdowns.")

    if label_regex:
        pat = re.compile(label_regex, re.I)
        # 1) casar pelo label do mapa
        for d in dropdowns:
            lab = (d.get("label") or "").strip()
            if lab and pat.search(lab):
                return d

        # 2) fallback por id conhecido no JSON do mapa
        #    (o id vem em dropdown["info"]["attrs"]["id"] quando o 00 capturou)
        def _id_of(d):
            return (((d.get("info") or {}).get("attrs") or {}).get("id")) or ""

        pref = normalize_text(label_regex)
        prefer_tipo = ("tipo" in pref)
        prefer_sub  = ("sub" in pref) or ("unidad" in pref) or ("secretar" in pref)

        order = []
        if prefer_tipo:
            order = ["slcTipo", "slcOrgsSubs", "slcOrgs"]
        elif prefer_sub:
            order = ["slcOrgsSubs", "slcOrgs", "slcTipo"]
        else:
            order = ["slcOrgs", "slcOrgsSubs", "slcTipo"]

        for wanted in order:
            for d in dropdowns:
                if _id_of(d) == wanted:
                    return d

    # 3) fallback por índice
    if len(dropdowns) > fallback_index:
        return dropdowns[fallback_index]
    raise RuntimeError(f"Dropdown de fallback index {fallback_index} não encontrado no mapa.")

def _filter_opts(
    options: List[Dict[str, Any]],
    select_regex: Optional[str],
    pick_list: Optional[str],
    limit: Optional[int]
) -> List[Dict[str, Any]]:

    
    opts = options or []
    out = opts

    if select_regex:
        # 1) Tentativa por regex padrão (respeitando acentos)
        try:
            pat = re.compile(select_regex, re.I)
            out = [o for o in opts if pat.search(o.get("text") or "")]
        except re.error:
            out = []

        # 2) Fallback por tokens (acentos-insensível), se a regex não retornou nada
        if not out:
            tokens = [t.strip() for t in select_regex.splitlines() if t.strip()]
            tokens_norm = [normalize_text(t) for t in tokens]
            tmp = []
            for o in opts:
                nt = normalize_text(o.get("text") or "")
                if any(tok and tok in nt for tok in tokens_norm):
                    tmp.append(o)
            out = tmp

    if pick_list:
        picks = {s.strip() for s in pick_list.split(",") if s.strip()}
        out = [o for o in out if (o.get("text") or "") in picks]

    if limit and limit > 0:
        out = out[:limit]

    return out

def _trim_placeholders(options: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    Remove entradas de placeholder típicas do DOU:
      - Selecionar Organização Principal/Subordinada/Tipo do Ato/Selecionar
      - Todos
    """
    bad = {
        "selecionar organizacao principal",
        "selecionar organização principal",
        "selecionar organizacao subordinada",
        "selecionar organização subordinada",
        "selecionar tipo do ato",
        "selecionar",
        "todos"
    }
    out = []
    for o in options or []:
        t = normalize_text(o.get("text") or "")
        if t in bad:
            continue
        out.append(o)
    return out

def _build_keys(opts: List[Dict[str, Any]], key_type: str) -> List[str]:
    keys = []
    for o in opts:
        if key_type == "text":
            t = (o.get("text") or "").strip()
            if t:
                keys.append(t)
        elif key_type == "value":
            v = o.get("value")
            if v not in (None, ""):
                keys.append(str(v))
        elif key_type == "dataValue":
            dv = o.get("dataValue")
            if dv not in (None, ""):
                keys.append(str(dv))
        elif key_type == "dataIndex":
            di = o.get("dataIndex")
            if di not in (None, ""):
                keys.append(str(di))
    seen = set(); out = []
    for k in keys:
        if k in seen: continue
        seen.add(k); out.append(k)
    return out

def plan_from_map(map_file: str, args) -> Dict[str, Any]:
    r"""
    Constrói o batch_config.json a partir do mapa (00), usando SOMENTE L1×L2.
    Melhorias:
      - Nível 3 sempre IGNORADO (força combos apenas L1×L2).
      - Remove placeholders ("Selecionar …", "Todos") antes de montar chaves.
      - Suporta --max-combos: corta o produto L1×L2 no limite desejado.
      - Diagnóstico (--plan-verbose): mostra roots/contagens/combos.
    """
    mp = json.loads(Path(map_file).read_text(encoding="utf-8"))
    scanned = mp.get("scannedUrl") or ""
    secao_from_map, data_from_map = _parse_secao_data_from_url(scanned)
    dropdowns = mp.get("dropdowns") or []
    if not dropdowns:
        raise RuntimeError("Mapa sem 'dropdowns'. Rode a 00 com --open-combos para capturar opções.")

    v = bool(getattr(args, "plan_verbose", False))

def plan_from_pairs(pairs_file: str, args) -> Dict[str, Any]:
    """
    Gera um batch_config.json a partir do artefato de mapeamento dinâmico N1->N2
    produzido pelo 00_map_pairs.py. Apenas combinações válidas (N1, N2 ∈ N1).

    Filtros suportados:
      - --select1 / --pick1 / --limit1           (aplicados sobre N1)
      - --select2 / --pick2 / --limit2-per-n1    (aplicados sobre N2 de cada N1)
      - --max-combos                             (corta o total gerado)
    Key types:
      - Escolha automática por opção (value > dataValue > id > dataId > dataIndex > text),
        a menos que o usuário force um default via --key1-type-default/--key2-type-default= text/value/...
    """
    pf = Path(pairs_file)
    data_pairs = json.loads(pf.read_text(encoding="utf-8"))

    data = data_pairs.get("date") or fmt_date(None)
    secao = data_pairs.get("secao") or getattr(args, "secao", "DO1")
    n1_groups = data_pairs.get("n1_options") or []

    # helpers locais usando os filtros já existentes
    def _apply_filter(options: List[Dict[str, Any]], sel: Optional[str], pick: Optional[str], lim: Optional[int]) -> List[Dict[str, Any]]:
        out = _filter_opts(options or [], sel, pick, lim)
        return out

    # 1) filtrar N1
    n1_filtered = []
    for grp in n1_groups:
        o1 = grp.get("n1") or {}
        n1_filtered.append(o1)

    n1_filtered = _apply_filter(n1_filtered, getattr(args, "select1", None), getattr(args, "pick1", None), getattr(args, "limit1", None))

    # Para lookup eficiente: mapeia N1 normalizado -> N2_options (do arquivo)
    def _key_norm(o): return normalize_text(o.get("text") or "")
    map_n1_to_n2 = { _key_norm(grp.get("n1") or {}): (grp.get("n2_options") or []) for grp in n1_groups }

    # 2) gerar combos válidos por N1
    combos: List[Dict[str, Any]] = []
    limit2_per_n1 = getattr(args, "limit2_per_n1", None)
    k1_def = getattr(args, "key1_type_default", None)  # se o usuário quiser forçar text/value/...
    k2_def = getattr(args, "key2_type_default", None)

    for o1 in n1_filtered:
        n1_norm = _key_norm(o1)
        n2_base = map_n1_to_n2.get(n1_norm, [])

        # aplica filtros em N2 deste N1
        n2_filtered = _apply_filter(
            n2_base,
            getattr(args, "select2", None),
            getattr(args, "pick2", None),
            limit2_per_n1  # limite por N1
        )

        # determina chaves para N1 (pode ser forçado ou automático)
        if k1_def:
            # usa o default do usuário
            if k1_def == "text":
                k1_type, k1_value = "text", (o1.get("text") or "").strip()
            elif k1_def == "value":
                k1_type, k1_value = "value", str(o1.get("value")) if o1.get("value") not in (None, "") else ( "text", (o1.get("text") or "").strip() )
                if isinstance(k1_value, tuple): k1_type, k1_value = k1_value  # fallback aplicado acima
            elif k1_def == "dataValue":
                dv = o1.get("dataValue")
                k1_type, k1_value = ("dataValue", str(dv)) if dv not in (None, "") else ("text", (o1.get("text") or "").strip())
            elif k1_def == "dataIndex":
                di = o1.get("dataIndex")
                k1_type, k1_value = ("dataIndex", str(di)) if di not in (None, "") else ("text", (o1.get("text") or "").strip())
            else:
                k1_type, k1_value = _best_key_for_option(o1)
        else:
            k1_type, k1_value = _best_key_for_option(o1)

        for o2 in n2_filtered:
            # determina chaves para N2
            if k2_def:
                if k2_def == "text":
                    k2_type, k2_value = "text", (o2.get("text") or "").strip()
                elif k2_def == "value":
                    k2_type, k2_value = "value", str(o2.get("value")) if o2.get("value") not in (None, "") else ("text", (o2.get("text") or "").strip())
                    if isinstance(k2_value, tuple): k2_type, k2_value = k2_value
                elif k2_def == "dataValue":
                    dv2 = o2.get("dataValue")
                    k2_type, k2_value = ("dataValue", str(dv2)) if dv2 not in (None, "") else ("text", (o2.get("text") or "").strip())
                elif k2_def == "dataIndex":
                    di2 = o2.get("dataIndex")
                    k2_type, k2_value = ("dataIndex", str(di2)) if di2 not in (None, "") else ("text", (o2.get("text") or "").strip())
                else:
                    k2_type, k2_value = _best_key_for_option(o2)
            else:
                k2_type, k2_value = _best_key_for_option(o2)

            combos.append({
                "key1_type": k1_type, "key1": k1_value,
                "key2_type": k2_type, "key2": k2_value,
                "key3_type": None, "key3": None,
                "label1": args.label1 or "", "label2": args.label2 or "", "label3": ""
            })

    # 3) cortar por --max-combos (se houver)
    maxc = getattr(args, "max_combos", None)
    if isinstance(maxc, int) and maxc > 0 and len(combos) > maxc:
        combos = combos[:maxc]

    if not combos:
        raise RuntimeError("Nenhum combo válido foi gerado a partir dos pares (verifique filtros/limites).")

    # 4) montar config final
    cfg = {
        "data": args.data or data,
        "secaoDefault": args.secao or secao or "DO1",
        "defaults": {
            "scrape_detail": bool(getattr(args, "scrape_detail", False)),
            "fallback_date_if_missing": bool(getattr(args, "fallback_date_if_missing", False)),
            "max_links": int(getattr(args, "max_links", 30)),
            "max_scrolls": int(getattr(args, "max_scrolls", 40)),
            "scroll_pause_ms": int(getattr(args, "scroll_pause_ms", 350)),
            "stable_rounds": int(getattr(args, "stable_rounds", 3)),
            "label1": args.label1, "label2": args.label2, "label3": None,
            "debug_dump": bool(getattr(args, "debug_dump", False)),
            "summary_lines": int(getattr(args, "summary_lines", 3)) if getattr(args, "summary_lines", None) else None,
            "summary_mode": getattr(args, "summary_mode", "center"),
        },
        "combos": combos,
        "output": {"pattern": "{secao}_{date}_{idx}.json", "report": "batch_report.json"}
    }

    # 5) topics (se houver --query)
    if getattr(args, "query", None):
        cfg["topics"] = [{"name": "Topic", "query": args.query}]

    if getattr(args, "state_file", None):
        cfg["state_file"] = args.state_file

    if getattr(args, "bulletin", None):
        ext = "docx" if args.bulletin == "docx" else args.bulletin
        out_b = args.bulletin_out or f"boletim_{{secao}}_{{date}}_{{idx}}.{ext}"
        cfg["output"]["bulletin"] = out_b
        cfg["defaults"]["bulletin"] = args.bulletin
        cfg["defaults"]["bulletin_out"] = out_b

    # combos não têm N3 por design
    for c in combos:
        c["key3_type"] = None
        c["key3"] = None

    return cfg


def plan_live(p, args) -> Dict[str, Any]:
    r"""
    Gera um plan DINÂMICO diretamente do site (somente L1×L2).
    Para cada N1 escolhido, seleciona o N1 na página, lê N2 ATUALIZADO e monta combos válidos.

    Suporta:
      - --select1/--pick1/--limit1  (pré-filtro de N1)
      - --select2/--pick2/--limit2  (filtro de N2 por N1)
      - --max-combos                (corta o total de combos)
      - --plan-verbose              (logs detalhados)

    Observação: N3 é IGNORADO por design (pedido do usuário).
    """
    def _trim_placeholders_local(options: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        bad = {
            "selecionar organizacao principal",
            "selecionar organização principal",
            "selecionar organizacao subordinada",
            "selecionar organização subordinada",
            "selecionar tipo do ato",
            "selecionar",
            "todos",
        }
        out = []
        for o in options or []:
            t = normalize_text(o.get("text") or "")
            if t in bad:
                continue
            out.append(o)
        return out

    v = bool(getattr(args, "plan_verbose", False))
    browser = p.chromium.launch(headless=not args.headful, slow_mo=args.slowmo)
    context = browser.new_context(ignore_https_errors=True, viewport={"width": 1366, "height": 900})
    page = context.new_page()
    page.set_default_timeout(60_000)
    page.set_default_navigation_timeout(60_000)

    combos: List[Dict[str, Any]] = []
    try:
        # 1) Abre página
        data = fmt_date(args.data)
        goto(page, data, args.secao)
        frame = find_best_frame(context)

        # 2) Encontra roots
        roots = collect_dropdown_roots(frame)
        if not roots:
            raise RuntimeError("Nenhum dropdown detectado.")
        r1 = resolve_dropdown(1, frame, page, roots, args.label1)
        r2 = resolve_dropdown(2, frame, page, roots, args.label2)
        if v:
            print("[plan-live] roots resolvidos (N1/N2).")

        # 3) Lê e filtra opções de N1
        o1 = read_dropdown_options(frame, r1)
        o1 = _trim_placeholders_local(o1)
        o1 = _filter_opts(o1, getattr(args, "select1", None), getattr(args, "pick1", None), getattr(args, "limit1", None))
        k1_list = _build_keys(o1, getattr(args, "key1_type_default", "text"))
        if not k1_list:
            raise RuntimeError("Após filtros, N1 ficou sem opções (ajuste --select1/--pick1/--limit1 e labels).")
        if v:
            print(f"[plan-live] N1 candidatos: {len(k1_list)}")

        # Config de corte
        maxc = getattr(args, "max_combos", None)
        if isinstance(maxc, int) and maxc <= 0:
            maxc = None

        # 4) Para cada N1, selecionar N1 e ler N2 atualizado
        for i, k1 in enumerate(k1_list, 1):
            # Se já atingiu o máximo de combos, pare
            if maxc and len(combos) >= maxc:
                break

            # (re)garante referências atuais de roots (DOM pode ter mudado)
            roots = collect_dropdown_roots(frame)
            r1 = resolve_dropdown(1, frame, page, roots, args.label1)

            if not r1 or not select_by_key(frame, r1, k1, getattr(args, "key1_type_default", "text")):
                if v: print(f"[plan-live][skip] N1 '{k1}' não pôde ser selecionado.")
                continue

            page.wait_for_load_state("networkidle", timeout=90_000)

            # Após selecionar N1, refaça roots e encontre N2
            roots = collect_dropdown_roots(frame)
            r2 = resolve_dropdown(2, frame, page, roots, args.label2)
            if not r2:
                if v: print(f"[plan-live][skip] N2 não encontrado após N1='{k1}'.")
                continue

            # Lê N2 atualizado e filtra
            o2 = read_dropdown_options(frame, r2)
            o2 = _trim_placeholders_local(o2)
            o2 = _filter_opts(o2, getattr(args, "select2", None), getattr(args, "pick2", None), getattr(args, "limit2", None))
            k2_list = _build_keys(o2, getattr(args, "key2_type_default", "text"))
            if v: print(f"[plan-live] N1='{k1}' => N2 válidos: {len(k2_list)}")
            if not k2_list:
                # Sem N2 válidos para este N1; siga para o próximo N1
                continue

            # 5) Monta combos válidos (somente L1×L2)
            for k2 in k2_list:
                combos.append({
                    "key1_type": getattr(args, "key1_type_default", "text"), "key1": k1,
                    "key2_type": getattr(args, "key2_type_default", "text"), "key2": k2,
                    "key3_type": None, "key3": None,                    # N3 OFF
                    "label1": r1.get("label") or "", "label2": r2.get("label") or "", "label3": ""
                })
                if maxc and len(combos) >= maxc:
                    break

        if v:
            print(f"[plan-live] combos gerados = {len(combos)}")

        if not combos:
            raise RuntimeError("Nenhum combo válido L1×L2 foi gerado (revise filtros e limites).")

        # 6) Monta config final (defaults herdados do CLI)
        cfg = {
            "data": data,
            "secaoDefault": args.secao or "DO1",
            "defaults": {
                "scrape_detail": bool(getattr(args, "scrape_detail", False)),
                "fallback_date_if_missing": bool(getattr(args, "fallback_date_if_missing", False)),
                "max_links": int(getattr(args, "max_links", 30)),
                "max_scrolls": int(getattr(args, "max_scrolls", 40)),
                "scroll_pause_ms": int(getattr(args, "scroll_pause_ms", 350)),
                "stable_rounds": int(getattr(args, "stable_rounds", 3)),
                "label1": args.label1, "label2": args.label2, "label3": None,
                "debug_dump": bool(getattr(args, "debug_dump", False)),
                "summary_lines": int(getattr(args, "summary_lines", 3)) if getattr(args, "summary_lines", None) else None,
                "summary_mode": getattr(args, "summary_mode", "center"),
            },
            "combos": combos,
            "output": {"pattern": "{secao}_{date}_{idx}.json", "report": "batch_report.json"}
        }

        if getattr(args, "query", None):
            cfg["topics"] = [{"name": "Topic", "query": args.query}]

        if getattr(args, "state_file", None):
            cfg["state_file"] = args.state_file

        if getattr(args, "bulletin", None):
            ext = "docx" if args.bulletin == "docx" else args.bulletin
            out_b = args.bulletin_out or f"boletim_{{secao}}_{{date}}_{{idx}}.{ext}"
            cfg["output"]["bulletin"] = out_b
            cfg["defaults"]["bulletin"] = args.bulletin
            cfg["defaults"]["bulletin_out"] = out_b

        # Garante combos sem N3
        for c in combos:
            c["key3_type"] = None
            c["key3"] = None

        return cfg

    finally:
        try: browser.close()
        except Exception: pass
    

    def _id_of(d):
        return (((d.get("info") or {}).get("attrs") or {}).get("id")) or ""


def _split_sentences(pt_text: str) -> List[str]:
    """
    Segmentação simples e robusta para PT-BR.
    Evita quebrar em abreviações comuns e normaliza espaços.
    """
    if not pt_text:
        return []
    text = pt_text

    # Protege abreviações comuns (Sr., Sra., Dr., Dra., et al.)
    text = re.sub(r"(?<=\bSr|Sra|Dr|Dra)\.", "", text)
    text = re.sub(r"(?<=\bet al)\.", "", text, flags=re.I)

    # Quebra por fim de frase . ! ?
    parts = re.split(r"(?<=[\.\!\?])\s+", text)
    sents = [re.sub(r"\s+", " ", s).strip() for s in parts if s and s.strip()]
    # Remove duplicadas consecutivas e linhas muito curtas "ocosas"
    cleaned = []
    seen = set()
    for s in sents:
        if len(s) < 20 and not any(ch.isalpha() for ch in s):
            continue
        key = s.lower()
        if key in seen:
            continue
        seen.add(key)
        cleaned.append(s)
    return cleaned

def _clean_text_for_summary(text: str) -> str:
    """
    Remove boilerplates comuns e normaliza espaços.
    Não é agressivo para evitar perda de conteúdo.
    """
    if not text:
        return ""
    t = re.sub(r"\s+", " ", text).strip()

    # Remover cabeçalhos/rodapés comuns
    patterns = [
        r"Este conteúdo não substitui.*?$",
        r"Publicado em\s*\d{1,2}/\d{1,2}/\d{2,4}.*?$",
        r"Imprensa Nacional.*?$",
        r"Ministério da.*?Diário Oficial.*?$",
    ]
    for pat in patterns:
        t = re.sub(pat, "", t, flags=re.I)

    # Remover repetições de título no início (rótulos)
    t = re.sub(r"^(?:\s*PORTARIA|RESOLUÇÃO|DECRETO|ATO|MENSAGEM)\s*[-–—]?\s*", "", t, flags=re.I)
    return t.strip()

def summarize_text(text: str,
                   max_lines: int = None,
                   keywords: Optional[List[str]] = None,
                   mode: str = "center") -> str:
    """
    Resumo por extração, focado no "centro da informação":
      - max_lines: nº de linhas (sentenças) desejadas.
      - mode:
         * center (default): equilíbrio entre densidade léxica, proximidade a keywords e posição central;
         * lead: primeiras sentenças (nariz de notícia);
         * keywords-first: prioriza sentenças que contêm keywords.
      - keywords: palavras de interesse (case-insensitive).
    """
    if not text:
        return ""

    max_lines = max_lines or SUMMARY_LINES or 7
    mode = mode or "center"

    # Preparação
    base = _clean_text_for_summary(text)
    sents = _split_sentences(base)
    if not sents:
        return ""


    # Se poucas sentenças, retorna "como está" dentro do limite
    if len(sents) <= max_lines:
        return "\n".join(sents[:max_lines])

    # Palavras-chave
    kws = [k.strip().lower() for k in (keywords or SUMMARY_KEYWORDS or []) if k.strip()]
    kw_set = set(kws)

    def tokens(s: str):
        raw = unicodedata.normalize("NFKD", s).encode("ascii", "ignore").decode("ascii")
        return [w for w in re.findall(r"[a-z0-9]+", raw.lower()) if len(w) >= 3]

    # Densidade léxica por sentença
    lex = []
    for s in sents:
        toks = tokens(s)
        unique = len(set(toks))
        lex.append(unique / (1 + len(toks)))  # favorece variedade, penaliza verborragia

    # Posição
    n = len(sents)
    pos = []
    for i in range(n):
        if mode == "lead":
            pos.append(1.0 - (i / n) * 0.8)  # início mais forte
        else:
            x = (i - (n - 1) / 2) / (n / 2)
            pos.append(1.0 - (x * x))  # pico no centro

    # Afinidade com keywords
    kscore = []
    for s in sents:
        st = s.lower()
        hits = sum(1 for k in kw_set if k and k in st)
        kscore.append(hits)

    scores = []
    for i, s in enumerate(sents):
        if mode == "keywords-first":
            w_k, w_l, w_p = 1.6, 1.0, 0.6
        elif mode == "lead":
            w_k, w_l, w_p = 1.0, 0.8, 1.6
        else:
            w_k, w_l, w_p = 1.4, 1.0, 1.2

        score = (w_k * kscore[i]) + (w_l * lex[i]) + (w_p * pos[i])

        # penalizações/ajustes
        if len(s) > 450:
            score -= 0.4
        if len(s) < 40 and kscore[i] == 0:
            score -= 0.5

        scores.append((score, i, s))

    scores.sort(key=lambda t: (-t[0], t[1]))
    picked_idx = sorted(i for _, i, _ in scores[:max_lines])

    # Completa lacunas se sobrar espaço
    if len(picked_idx) < max_lines:
        need = max_lines - len(picked_idx)
        pool = [i for i in range(n) if i not in picked_idx]
        anchor = scores[0][1]
        pool.sort(key=lambda i: abs(i - anchor))
        picked_idx.extend(pool[:need])
        picked_idx = sorted(set(picked_idx))[:max_lines]

    out_lines = [sents[i].strip() for i in picked_idx]

    # Dedup leve e poda final
    final, seen = [], set()
    for ln in out_lines:
        key = ln.lower()
        if key in seen:
            continue
        seen.add(key)
        final.append(ln)
        if len(final) >= max_lines:
            break
    return "\n".join(final[:max_lines]).strip()

# ------------------------- Batch / Presets -------------------------
def expand_batch_config(cfg: Dict[str, Any]) -> List[Dict[str, Any]]:
    """
    Expande o config em uma lista de jobs:
      - se houver cfg["jobs"], usa-os e aplica defaults;
      - se houver topics e combos, faz o produto cartesiano topics × combos;
      - se NÃO houver topics, mas houver combos, gera jobs "combos-only".
    PATCH: propaga summary_keywords/summary_lines/summary_mode do tópico para o job.
    """
    jobs: List[Dict[str, Any]] = []
    defaults = cfg.get("defaults", {})
    base_data = cfg.get("data")
    base_secao = cfg.get("secaoDefault")

    def merge_defaults(job):
        out = dict(defaults)
        out.update(job or {})
        return out

    # 1) Jobs explícitos
    if cfg.get("jobs"):
        for j in cfg["jobs"]:
            jj = merge_defaults(j)
            if base_data and not jj.get("data"):
                jj["data"] = base_data
            if base_secao and not jj.get("secao"):
                jj["secao"] = base_secao
            rep = max(1, int(jj.get("repeat", cfg.get("repeat", 1))))
            for r in range(1, rep + 1):
                jj_r = dict(jj); jj_r["_repeat"] = r
                jobs.append(jj_r)

    # 2) topics × combos
    topics = cfg.get("topics") or []
    combos = cfg.get("combos") or []
    if topics and combos:
        for t in topics:
            topic_name = t.get("name") or "topic"
            topic_query = t.get("query") or ""
            topic_repeat = int(t.get("repeat", cfg.get("repeat", 1)))
            # PATCH: pegar overrides de resumo no tópico
            t_summary_kws  = t.get("summary_keywords")
            t_summary_lines = t.get("summary_lines", (defaults.get("summary_lines") if isinstance(defaults, dict) else None))
            t_summary_mode  = t.get("summary_mode", (defaults.get("summary_mode") if isinstance(defaults, dict) else None))

            for idx, c in enumerate(combos, 1):
                jj = merge_defaults(c)
                if base_data and not jj.get("data"):
                    jj["data"] = base_data
                if base_secao and not jj.get("secao"):
                    jj["secao"] = base_secao
                jj["topic"] = topic_name
                jj["query"] = jj.get("query", topic_query)
                jj["_combo_index"] = idx

                # PATCH: acoplar overrides de resumo no job
                if t_summary_kws is not None:
                    jj["summary_keywords"] = t_summary_kws
                if t_summary_lines is not None:
                    jj["summary_lines"] = t_summary_lines
                if t_summary_mode is not None:
                    jj["summary_mode"] = t_summary_mode

                rep = max(1, int(jj.get("repeat", topic_repeat)))
                for r in range(1, rep + 1):
                    jj_r = dict(jj); jj_r["_repeat"] = r
                    jobs.append(jj_r)

    # 3) combos-only
    if not topics and combos:
        for idx, c in enumerate(combos, 1):
            jj = merge_defaults(c)
            if base_data and not jj.get("data"):
                jj["data"] = base_data
            if base_secao and not jj.get("secao"):
                jj["secao"] = base_secao
            jj["topic"] = jj.get("topic") or f"job{idx}"
            jj["query"] = jj.get("query", defaults.get("query", "") if isinstance(defaults, dict) else "")
            jj["_combo_index"] = idx
            rep = max(1, int(jj.get("repeat", cfg.get("repeat", 1))))
            for r in range(1, rep + 1):
                jj_r = dict(jj); jj_r["_repeat"] = r
                jobs.append(jj_r)

    return jobs

def render_out_filename(pattern: str, job: Dict[str, Any]) -> str:
    date_str = job.get("data") or fmt_date(None)
    tokens = {
        "topic": job.get("topic") or "job",
        "secao": job.get("secao") or "DO",
        "date": date_str.replace("/", "-"),
        "idx": str(job.get("_combo_index") or job.get("_job_index") or ""),
        "rep": str(job.get("_repeat") or ""),
        "key1": job.get("key1") or "",
        "key2": job.get("key2") or "",
        "key3": job.get("key3") or "",
    }
    name = pattern
    for k, v in tokens.items():
        name = name.replace("{" + k + "}", sanitize_filename(str(v)))
    return name

def iter_dates(args_data: Optional[str], since: int, range_arg: Optional[str]):
    if range_arg:
        a, b = [s.strip() for s in range_arg.split(":")]
        d0 = datetime.strptime(a, "%d-%m-%Y").date()
        d1 = datetime.strptime(b, "%d-%m-%Y").date()
        step = timedelta(days=1)
        cur = d0
        while cur <= d1:
            yield cur.strftime("%d-%m-%Y")
            cur += step
        return
    base = datetime.strptime(fmt_date(args_data), "%d-%m-%Y").date()
    for off in range(0, max(0, since) + 1):
        d = base - timedelta(days=off)
        yield d.strftime("%d-%m-%Y")

def reset_dropdown(frame, root: Dict[str, Any]) -> None:
    """
    Tenta resetar um dropdown para o primeiro item (index 0).
    Somente <select> nativo; em custom combobox, um reset genérico não é confiável.
    """
    try:
        if is_select_root(root):
            sel = root["handle"] if isinstance(root, dict) else root
            try:
                sel.select_option(index=0)
                frame.page.wait_for_load_state("networkidle", timeout=60_000)
            except Exception:
                pass
    except Exception:
        pass


def flow_batch(p, args):
    """
    Executa um plano (config JSON) em modo batch, com melhorias:
      - Reuso opcional de aba por (data, secao);
      - Dedup global por state_file carregado uma única vez;
      - Reset de N1 antes de cada seleção para reduzir falhas por estado residual;
      - Debug dump automático em falhas (quando --debug-dump);
      - Boletim por job opcional.
    """
    from collections import defaultdict
    import hashlib

    # --- carregar config e preparar saída ---
    cfg_path = Path(args.config)
    cfg = json.loads(cfg_path.read_text(encoding="utf-8"))
    out_dir = Path(args.out_dir or ".")
    out_dir.mkdir(parents=True, exist_ok=True)

    # Expande jobs a partir do config (topics × combos ou combos-only)
    jobs = expand_batch_config(cfg)
    if not jobs:
        print("[Erro] Nenhum job gerado a partir do config.")
        sys.exit(1)

    out_pattern = (cfg.get("output") or {}).get("pattern") or "{topic}_{secao}_{date}_{idx}.json"
    report = {"total_jobs": len(jobs), "ok": 0, "fail": 0, "items_total": 0, "outputs": []}

    # Defaults do config
    dfl = cfg.get("defaults") or {}

    def _get(job, key, default_key=None, default_value=None):
        """Busca chave do job com fallback nos defaults do config."""
        if default_key is None:
            default_key = key
        return job.get(key, dfl.get(default_key, default_value))

    # --- Dedup global: carrega state_file (se existir) apenas 1 vez ---
    state_file_path = None
    global_seen = set()
    if cfg.get("state_file"):
        state_file_path = Path(cfg["state_file"])
    elif getattr(args, "state_file", None):
        state_file_path = Path(args.state_file)
    if state_file_path and state_file_path.exists():
        try:
            for line in state_file_path.read_text(encoding="utf-8").splitlines():
                try:
                    obj = json.loads(line); h = obj.get("hash")
                    if h: global_seen.add(h)
                except Exception:
                    pass
        except Exception:
            pass

    # --- Browser/context único para todo o batch ---
    browser = p.chromium.launch(headless=not args.headful, slow_mo=args.slowmo)
    context = browser.new_context(ignore_https_errors=True, viewport={"width": 1366, "height": 900})

    try:
        if getattr(args, "reuse_page", False):
            # ===== RAMO COM REUSO DE ABA POR (data, secao) =====
            groups = defaultdict(list)
            for global_idx, job in enumerate(jobs, 1):
                jd = fmt_date(job.get("data"))
                js = job.get("secao", cfg.get("secaoDefault", "DO1"))
                groups[(jd, js)].append((global_idx, job))

            for (data, secao), gjobs in groups.items():
                print(f"\n[Grupo] {data} / {secao} — {len(gjobs)} jobs")

                page = context.new_page()
                page.set_default_timeout(60_000)
                page.set_default_navigation_timeout(60_000)

                try:
                    # 1) Navegação única por (data, secao)
                    goto(page, data, secao)
                    frame = find_best_frame(context)

                    # 2) Processa todos os jobs do grupo
                    for global_idx, job in gjobs:
                        topic = job.get("topic", "")
                        try:
                            # Parâmetros do job
                            key1_type = job.get("key1_type"); key1 = job.get("key1")
                            key2_type = job.get("key2_type"); key2 = job.get("key2")
                            if not key1 or not key1_type or not key2 or not key2_type:
                                raise ValueError("Faltando key1/key1_type ou key2/key2_type")

                            max_links     = int(_get(job, "max_links", "max_links", 30))
                            max_scrolls   = int(_get(job, "max_scrolls", "max_scrolls", 40))
                            scroll_pause  = int(_get(job, "scroll_pause_ms", "scroll_pause_ms", 350))
                            stable_rounds = int(_get(job, "stable_rounds", "stable_rounds", 3))
                            debug_dump    = bool(_get(job, "debug_dump", "debug_dump", False))
                            scrape_detail = bool(_get(job, "scrape_detail", "scrape_detail", False))
                            detail_timeout = int(_get(job, "detail_timeout", "detail_timeout", 60_000))
                            fallback_date  = bool(_get(job, "fallback_date_if_missing", "fallback_date_if_missing", True))

                            state_file_local = job.get("state_file") or cfg.get("state_file")
                            bulletin   = job.get("bulletin") or dfl.get("bulletin")
                            bulletin_out_pat = job.get("bulletin_out") or dfl.get("bulletin_out") or (cfg.get("output") or {}).get("bulletin")

                            # nome do JSON de saída (usa índice global)
                            out_name = render_out_filename(out_pattern, {**job, "_job_index": global_idx})
                            out_path = out_dir / out_name
                            out_path.parent.mkdir(parents=True, exist_ok=True)

                            # Ajusta overrides de resumo no job (se existirem)
                            _apply_summary_overrides_from_job(job)

                            # --- Seleciona N1 ---
                            roots = collect_dropdown_roots(frame)
                            r1 = resolve_dropdown(1, frame, page, roots, job.get("label1"))
                            if r1:
                                reset_dropdown(frame, r1)
                            if not r1 or not select_by_key(frame, r1, key1, key1_type):
                                print(f"[Info] N1 falhou: {key1_type}='{key1}' — {topic}")
                                continue
                            page.wait_for_load_state("networkidle", timeout=90_000)

                            # --- Seleciona N2 ---
                            roots = collect_dropdown_roots(frame)
                            r2 = resolve_dropdown(2, frame, page, roots, job.get("label2"))
                            if not r2 or not select_by_key(frame, r2, key2, key2_type):
                                print(f"[Info] N2 falhou: {key2_type}='{key2}' — {topic}")
                                continue
                            page.wait_for_load_state("networkidle", timeout=90_000)

                            # --- Query + coleta ---
                            query = job.get("query", "")
                            apply_query(frame, query)
                            itens = collect_links(
                                frame, max_links,
                                max_scrolls=max_scrolls,
                                scroll_pause_ms=scroll_pause,
                                stable_rounds=stable_rounds
                            )

                            # --- Enriquecimento (opcional) ---
                            enriched = itens
                            if scrape_detail:
                                print(f"[Info] Enriquecendo {len(itens)} itens (timeout={detail_timeout} ms cada)...")
                                enriched = []
                                for it in itens:
                                    detail_url = abs_url(page.url, it.get("link"))
                                    try:
                                        meta = scrape_detail(context, detail_url, timeout_ms=detail_timeout)
                                    except Exception:
                                        meta = {"detail_url": detail_url}

                                    def _is_bad(t: str) -> bool:
                                        return normalize_text(t) in {"imprensa nacional", ""}

                                    titulo_list = it.get("titulo") or ""
                                    titulo_det  = (meta.get("titulo") or "").strip()
                                    final_title = titulo_det if not _is_bad(titulo_det) else titulo_list

                                    data_pub = meta.get("data_publicacao")
                                    if not data_pub and fallback_date:
                                        try:
                                            dt = datetime.strptime(data, "%d-%m-%Y").date()
                                            data_pub = dt.strftime("%Y-%m-%d")
                                        except Exception:
                                            pass

                                    # atribui hash (detalhe/link) para dedup
                                    lk = detail_url or it.get("link") or ""
                                    h = hashlib.sha1(lk.encode("utf-8")).hexdigest() if lk else None

                                    enriched.append({
                                        **it, **meta,
                                        "titulo_listagem": titulo_list,
                                        "titulo_detalhe":  titulo_det,
                                        "titulo": final_title,
                                        "detail_url": detail_url,
                                        "data_publicacao": data_pub or meta.get("data_publicacao"),
                                        "data_publicacao_fallback": (not meta.get("data_publicacao")),
                                        "hash": h,
                                    })

                            # --- Dedup global via state_file ---
                            filtered = []
                            if state_file_path:
                                for it in enriched:
                                    h = it.get("hash")
                                    if not h:
                                        # fallback: faz hash do link/detalhe se ainda não houver
                                        lk = it.get("detail_url") or it.get("link") or ""
                                        h = hashlib.sha1(lk.encode("utf-8")).hexdigest() if lk else None
                                        if h:
                                            it["hash"] = h
                                    if h and h in global_seen:
                                        continue
                                    filtered.append(it)
                                    if h:
                                        global_seen.add(h)
                                # atualiza state_file incrementalmente
                                try:
                                    if filtered:
                                        state_file_path.parent.mkdir(parents=True, exist_ok=True)
                                        with state_file_path.open("a", encoding="utf-8") as f:
                                            for it in filtered:
                                                if it.get("hash"):
                                                    f.write(json.dumps({"hash": it["hash"]}, ensure_ascii=False) + "\n")
                                except Exception:
                                    pass
                            else:
                                filtered = enriched

                            # --- Salvar JSON do job ---
                            result = {
                                "data": data,
                                "secao": secao,
                                "selecoes": [
                                    {"level": 1, "type": key1_type, "key": key1},
                                    {"level": 2, "type": key2_type, "key": key2},
                                ],
                                "query": job.get("query", ""),
                                "total": len(filtered),
                                "itens": filtered,
                                "enriquecido": bool(scrape_detail),
                            }
                            out_path.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
                            print(f"[OK] {out_path}")
                            report["ok"] += 1
                            report["outputs"].append(str(out_path))
                            report["items_total"] += result["total"]

                            # --- Boletim por job (opcional) ---
                            if bulletin and bulletin_out_pat:
                                b_out = out_dir / render_out_filename(bulletin_out_pat, {**job, "_job_index": global_idx})
                                b_out.parent.mkdir(parents=True, exist_ok=True)
                                try:
                                    gen_bulletin(result, bulletin, str(b_out))
                                    print(f"[OK] Boletim gerado: {b_out}")
                                except Exception as e:
                                    print(f"[Aviso] Falha ao gerar boletim: {e}")

                        except Exception as e:
                            print(f"[FAIL] Job (grupo {data}/{secao}) idx={global_idx}: {e}")
                            if bool(_get(job, "debug_dump", "debug_dump", False)):
                                dump_debug(page, f"debug_job_group_{data}_{secao}_{global_idx}")
                            report["fail"] += 1

                finally:
                    try:
                        page.close()
                    except Exception:
                        pass

        else:
            # ===== RAMO SEM REUSO (um page por job) =====
            for j_idx, job in enumerate(jobs, 1):
                print(f"\n[Job {j_idx}/{len(jobs)}] {job.get('topic','')}: {job.get('query','')}")
                out_name = render_out_filename(out_pattern, {**job, "_job_index": j_idx})
                out_path = out_dir / out_name

                data = fmt_date(job.get("data"))
                secao = job.get("secao", cfg.get("secaoDefault", "DO1"))
                key1_type = job.get("key1_type"); key1 = job.get("key1")
                key2_type = job.get("key2_type"); key2 = job.get("key2")
                label1 = job.get("label1"); label2 = job.get("label2")

                max_links = int(_get(job, "max_links", "max_links", 30))
                scrape_detail = bool(_get(job, "scrape_detail", "scrape_detail", True))
                detail_timeout = int(_get(job, "detail_timeout", "detail_timeout", 60_000))
                fallback_date = bool(_get(job, "fallback_date_if_missing", "fallback_date_if_missing", True))
                debug_dump = bool(_get(job, "debug_dump", "debug_dump", False))

                max_scrolls = int(_get(job, "max_scrolls", "max_scrolls", 40))
                scroll_pause_ms = int(_get(job, "scroll_pause_ms", "scroll_pause_ms", 350))
                stable_rounds = int(_get(job, "stable_rounds", "stable_rounds", 3))

                bulletin = job.get("bulletin") or dfl.get("bulletin")
                bulletin_out_pat = job.get("bulletin_out") or dfl.get("bulletin_out") or (cfg.get("output") or {}).get("bulletin")
                bulletin_out = out_dir / render_out_filename(bulletin_out_pat, {**job, "_job_index": j_idx}) if (bulletin and bulletin_out_pat) else None

                page = context.new_page()
                page.set_default_timeout(60_000)
                page.set_default_navigation_timeout(60_000)

                try:
                    missing = []
                    if not key1 or not key1_type: missing.append("key1/key1_type")
                    if not key2 or not key2_type: missing.append("key2/key2_type")
                    if missing:
                        raise ValueError(f"Parâmetros faltando: {', '.join(missing)}")

                    # Overrides de resumo para este job (se houverem)
                    _apply_summary_overrides_from_job(job)

                    # Fluxo inline (equivalente ao flow_run, mantendo dependências locais)
                    goto(page, data, secao)
                    frame = find_best_frame(context)

                    roots = collect_dropdown_roots(frame)
                    r1 = resolve_dropdown(1, frame, page, roots, label1)
                    if r1:
                        reset_dropdown(frame, r1)
                    if not r1 or not select_by_key(frame, r1, key1, key1_type):
                        raise RuntimeError(f"N1 falhou: {key1_type}='{key1}'")
                    page.wait_for_load_state("networkidle", timeout=90_000)

                    roots = collect_dropdown_roots(frame)
                    r2 = resolve_dropdown(2, frame, page, roots, label2)
                    if not r2 or not select_by_key(frame, r2, key2, key2_type):
                        raise RuntimeError(f"N2 falhou: {key2_type}='{key2}'")
                    page.wait_for_load_state("networkidle", timeout=90_000)

                    query = job.get("query", "")
                    apply_query(frame, query)
                    itens = collect_links(
                        frame, max_links,
                        max_scrolls=max_scrolls,
                        scroll_pause_ms=scroll_pause_ms,
                        stable_rounds=stable_rounds
                    )

                    enriched = itens
                    if scrape_detail:
                        print(f"[Info] Enriquecendo {len(itens)} itens (timeout={detail_timeout} ms cada)...")
                        enriched = []
                        for it in itens:
                            detail_url = abs_url(page.url, it.get("link"))
                            try:
                                meta = scrape_detail(context, detail_url, timeout_ms=detail_timeout)
                            except Exception:
                                meta = {"detail_url": detail_url}

                            def _is_bad(t: str) -> bool:
                                return normalize_text(t) in {"imprensa nacional", ""}

                            titulo_list = it.get("titulo") or ""
                            titulo_det  = (meta.get("titulo") or "").strip()
                            final_title = titulo_det if not _is_bad(titulo_det) else titulo_list

                            data_pub = meta.get("data_publicacao")
                            if not data_pub and fallback_date:
                                try:
                                    dt = datetime.strptime(data, "%d-%m-%Y").date()
                                    data_pub = dt.strftime("%Y-%m-%d")
                                except Exception:
                                    pass

                            lk = detail_url or it.get("link") or ""
                            h = hashlib.sha1(lk.encode("utf-8")).hexdigest() if lk else None

                            enriched.append({
                                **it, **meta,
                                "titulo_listagem": titulo_list,
                                "titulo_detalhe":  titulo_det,
                                "titulo": final_title,
                                "detail_url": detail_url,
                                "data_publicacao": data_pub or meta.get("data_publicacao"),
                                "data_publicacao_fallback": (not meta.get("data_publicacao")),
                                "hash": h,
                            })

                    # Dedup global
                    filtered = []
                    if state_file_path:
                        for it in enriched:
                            h = it.get("hash")
                            if not h:
                                lk = it.get("detail_url") or it.get("link") or ""
                                h = hashlib.sha1(lk.encode("utf-8")).hexdigest() if lk else None
                                if h:
                                    it["hash"] = h
                            if h and h in global_seen:
                                continue
                            filtered.append(it)
                            if h:
                                global_seen.add(h)
                        try:
                            if filtered:
                                state_file_path.parent.mkdir(parents=True, exist_ok=True)
                                with state_file_path.open("a", encoding="utf-8") as f:
                                    for it in filtered:
                                        if it.get("hash"):
                                            f.write(json.dumps({"hash": it["hash"]}, ensure_ascii=False) + "\n")
                        except Exception:
                            pass
                    else:
                        filtered = enriched

                    # Salvar JSON do job
                    result = {
                        "data": data,
                        "secao": secao,
                        "selecoes": [
                            {"level": 1, "type": key1_type, "key": key1},
                            {"level": 2, "type": key2_type, "key": key2},
                        ],
                        "query": query,
                        "total": len(filtered),
                        "itens": filtered,
                        "enriquecido": bool(scrape_detail),
                    }
                    out_path.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
                    print(f"[OK] {out_path}")
                    report["ok"] += 1
                    report["outputs"].append(str(out_path))
                    report["items_total"] += result["total"]

                    if bulletin and bulletin_out:
                        try:
                            gen_bulletin(result, bulletin, str(bulletin_out))
                            print(f"[OK] Boletim gerado: {bulletin_out}")
                        except Exception as e:
                            print(f"[Aviso] Falha ao gerar boletim: {e}")

                except Exception as e:
                    print(f"[FAIL] Job {j_idx}: {e}")
                    if debug_dump:
                        dump_debug(page, f"debug_job_{j_idx}")
                    report["fail"] += 1
                finally:
                    try:
                        page.close()
                    except Exception:
                        pass

                delay_ms = int(job.get("repeat_delay_ms", cfg.get("repeat_delay_ms", 0)))
                if delay_ms > 0:
                    time.sleep(delay_ms / 1000.0)

    finally:
        try:
            browser.close()
        except Exception:
            pass

    rep_path = out_dir / (((cfg.get("output", {}) or {}).get("report")) or "batch_report.json")
    rep_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"\n[REPORT] {rep_path} — jobs={report['total_jobs']} ok={report['ok']} fail={report['fail']} items={report['items_total']}")

def flow_report(in_dir: str, kind: str, out_path: str, date_label: str = None, secao_label: str = None):
    """
    Consolida todos os JSONs de 'in_dir' e gera um único boletim (docx/md/html).
    Usa gen_bulletin(result, kind, out_path).
    """
    from pathlib import Path
    import json as _json

    # garantir pasta do boletim final
    Path(out_path).parent.mkdir(parents=True, exist_ok=True)

    agg = []
    for f in sorted(Path(in_dir).glob("*.json")):
        try:
            data = _json.loads(f.read_text(encoding="utf-8"))
            agg.extend(data.get("itens", []))
        except Exception:
            pass

    result = {
        "data": date_label or "",
        "secao": secao_label or "",
        "total": len(agg),
        "itens": agg
    }
    gen_bulletin(result, kind, out_path)
    print(f"[OK] Boletim consolidado gerado: {out_path}")


# ------------------------- Presets -------------------------
def load_presets(path: Path) -> Dict[str, Any]:
    if path.exists():
        try:
            return json.loads(path.read_text(encoding="utf-8"))
        except Exception:
            pass
    return {"presets": {}}

def save_presets(path: Path, data: Dict[str, Any]) -> None:
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")

def build_preset_from_args(args) -> Dict[str, Any]:
    return {
        "mode": "run",
        "data": args.data,
        "secao": args.secao,
        "key1_type": args.key1_type, "key1": args.key1,
        "key2_type": args.key2_type, "key2": args.key2,
        "key3_type": args.key3_type, "key3": args.key3,
        "label1": args.label1, "label2": args.label2, "label3": args.label3,
        "query": args.query,
        "max_links": args.max_links,
        "max_scrolls": args.max_scrolls,
        "scroll_pause_ms": args.scroll_pause_ms,
        "stable_rounds": args.stable_rounds,
        "scrape_detail": args.scrape_detail,
        "detail_timeout": args.detail_timeout,
        "fallback_date_if_missing": args.fallback_date_if_missing if hasattr(args, 'fallback_date_if_missing') else False,
        "debug_dump": args.debug_dump,
        "state_file": args.state_file,
        "bulletin": args.bulletin,
        "bulletin_out": args.bulletin_out,
    }

# ------------------------- Main CLI -------------------------
def main():
    ap = argparse.ArgumentParser(description="Cascata DOU v6.1 (00->plan->batch/run + boletim + dedup)")
    ap.add_argument("--mode", choices=["list", "run", "batch", "plan", "report", "plan-live", "plan-from-pairs"], required=True)
    ap.add_argument("--level", type=int, default=1, help="Para list: 1\n2\n3")

    # filtros
    ap.add_argument("--key1-type", choices=["value", "dataValue", "dataIndex", "text"])
    ap.add_argument("--key1")
    ap.add_argument("--key2-type", choices=["value", "dataValue", "dataIndex", "text"])
    ap.add_argument("--key2")
    ap.add_argument("--key3-type", choices=["value", "dataValue", "dataIndex", "text"])
    ap.add_argument("--key3")

    # rótulos
    ap.add_argument("--label1", help="Rótulo do nível 1 (ex.: 'Órgão')")
    ap.add_argument("--label2", help="Rótulo do nível 2 (ex.: 'Secretaria\nUnidade')")
    ap.add_argument("--label3", help="Rótulo do nível 3 (ex.: 'Tipo do Ato')")

    # gerais
    ap.add_argument("--query", default=None)
    ap.add_argument("--max-links", type=int, default=30)
    ap.add_argument("--max-scrolls", type=int, default=40, help="Máximo de scrolls para carregar resultados")
    ap.add_argument("--scroll-pause-ms", type=int, default=350, help="Pausa entre scrolls (ms)")
    ap.add_argument("--stable-rounds", type=int, default=3, help="Rodadas de estabilidade para parar scroll")
    ap.add_argument("--secao", default="DO1")
    ap.add_argument("--data", default=None, help="DD-MM-AAAA (default: hoje)")
    ap.add_argument("--since", type=int, default=0, help="Executa também N dias para trás (0 = apenas a data)")
    ap.add_argument("--range", help="Faixa de datas: DD-MM-AAAA:DD-MM-AAAA")
    ap.add_argument("--out", default="out_{date}.json")

    # detalhes
    ap.add_argument("--scrape-detail", action="store_true", help="Abrir cada link e extrair metadados do detalhe")
    ap.add_argument("--detail-timeout", type=int, default=60_000, help="Timeout (ms) por detalhe")
    ap.add_argument("--fallback-date-if-missing", action="store_true",
                    help="Se data_publicacao não for encontrada, usa a data da edição (--data)")

    # dedup e boletim
    ap.add_argument("--state-file", help="JSONL de estado para deduplicação por hash")
    ap.add_argument("--bulletin", choices=["html", "md", "docx"], help="Gera boletim (html/md/docx) agrupado por órgão/tipo")
    ap.add_argument("--bulletin-out", help="Arquivo de saída do boletim (ex.: boletim_ri_{date}.html)")

    # batch
    ap.add_argument("--config", help="Arquivo JSON de configuração (mode=batch)")
    ap.add_argument("--out-dir", help="Diretório base de saída para batch")
    ap.add_argument("--reuse-page", action="store_true", help="(batch) Reutiliza uma mesma aba por data/seção para todos os jobs")

    # presets
    ap.add_argument("--save-preset", help="Nome do preset a salvar (dos args atuais de run)")
    ap.add_argument("--preset", help="Nome do preset a carregar e executar (mode=run)")
    ap.add_argument("--presets-file", default="presets.json", help="Arquivo de presets (default: presets.json)")
    ap.add_argument("--save-preset-no-date", action="store_true", help="Salva preset sem a data (usa 'hoje' ao executar)")

    # PLAN
    ap.add_argument("--map-file", help="JSON gerado pela rotina 00 (mapa da página com opções dos dropdowns)")
    ap.add_argument("--plan-out", default="batch_config.json", help="Arquivo de saída do plano/batch gerado")
    ap.add_argument("--execute-plan", action="store_true", help="Após gerar o plano, executar o batch automaticamente")
    ap.add_argument("--select1", help="Regex para filtrar opções do nível 1 (por texto)")
    ap.add_argument("--select2", help="Regex para filtrar opções do nível 2 (por texto)")
    ap.add_argument("--select3", help="Regex para filtrar opções do nível 3 (por texto)")
    ap.add_argument("--pick1", help="Lista fixa (vírgula) de opções do nível 1 (match por texto)")
    ap.add_argument("--pick2", help="Lista fixa (vírgula) de opções do nível 2 (match por texto)")
    ap.add_argument("--pick3", help="Lista fixa (vírgula) de opções do nível 3 (match por texto)")
    ap.add_argument("--limit1", type=int, help="Limite de opções no nível 1")
    ap.add_argument("--limit2", type=int, help="Limite de opções no nível 2")
    ap.add_argument("--limit3", type=int, help="Limite de opções no nível 3")
    ap.add_argument("--key1-type-default", choices=["text","value","dataValue","dataIndex"], default="text")
    ap.add_argument("--key2-type-default", choices=["text","value","dataValue","dataIndex"], default="text")
    ap.add_argument("--key3-type-default", choices=["text","value","dataValue","dataIndex"], default="text")
    ap.add_argument("--no-level3", action="store_true", help="(obsoleto no plan: N3 é ignorado)")
    ap.add_argument("--plan-verbose", action="store_true", help="Mostra diagnóstico detalhado durante o plan")
    ap.add_argument("--max-combos", type=int, help="(plan) Corta o produto L1×L2 para no máximo N")  # IMPORTANTE: defina só UMA vez!

    # PLAN from pairs
    ap.add_argument("--pairs-file", help="JSON de pares N1->N2 gerado pelo 00_map_pairs.py")
    ap.add_argument("--limit2-per-n1", type=int, help="Limite de N2 por N1 (ao gerar combos a partir dos pares)")

    # ambiente
    ap.add_argument("--headful", action="store_true")
    ap.add_argument("--slowmo", type=int, default=0)
    ap.add_argument("--debug-dump", action="store_true", help="Salvar screenshot/DOM em caso de erro")

    # Resumo (globais)
    ap.add_argument("--summary-sentences", type=int, default=3, help="Nº de frases no resumo (default: 3)")
    ap.add_argument("--summary-keywords", help="Lista de palavras de interesse separadas por ponto e vírgula (;)")
    ap.add_argument("--summary-keywords-file", help="Arquivo .txt com palavras de interesse (uma por linha)")
    ap.add_argument("--summary-lines", type=int, help="Nº de linhas do resumo (preferido). Se não for passado, usa --summary-sentences.")
    ap.add_argument("--summary-mode", choices=["center", "lead", "keywords-first"], default="center",
                    help="Estratégia de resumo: center (padrão), lead ou keywords-first")

    # REPORT
    ap.add_argument("--in-dir", help="(report) Pasta com os JSONs (saídas dos jobs)")
    ap.add_argument("--report-out", help="(report) Caminho do boletim final (docx/md/html)")

    args = ap.parse_args()
    _setup_summary_globals(args)

    # Salvar preset (opcional)
    if args.save_preset:
        preset_path = Path(args.presets_file)
        store = load_presets(preset_path)
        if not hasattr(args, "fallback_date_if_missing"):
            args.fallback_date_if_missing = False
        pr = build_preset_from_args(args)
        if args.save_preset_no_date:
            pr["data"] = None
        store["presets"][args.save_preset] = pr
        save_presets(preset_path, store)
        print(f"[OK] Preset salvo: {args.save_preset} em {preset_path}")

    # REPORT
    if args.mode == "report":
        if not args.in_dir or not args.report_out or not args.bulletin:
            print("[Erro] --in-dir, --report-out e --bulletin são obrigatórios no mode=report"); sys.exit(1)
        flow_report(args.in_dir, args.bulletin, args.report_out, date_label=args.data, secao_label=args.secao)
        return

    # plan-from-pairs (não precisa Playwright; só se --execute-plan)
    if args.mode == "plan-from-pairs":
        if not args.pairs_file:
            print("[Erro] --pairs-file obrigatório em mode=plan-from-pairs"); sys.exit(1)
        try:
            cfg = plan_from_pairs(args.pairs_file, args)
            Path(args.plan_out).parent.mkdir(parents=True, exist_ok=True)
            Path(args.plan_out).write_text(json.dumps(cfg, ensure_ascii=False, indent=2), encoding="utf-8")
            print(f"[OK] Plano (from-pairs) gerado em: {args.plan_out} (combos={len(cfg.get('combos', []))})")
            if args.execute_plan:
                print("[Exec] Rodando batch com o plano recém-gerado...")
                args.config = args.plan_out
                from playwright.sync_api import sync_playwright
                with sync_playwright() as p:
                    flow_batch(p, args)
        except Exception as e:
            print(f"[ERRO][plan-from-pairs] {e}")
            if getattr(args, "plan_verbose", False):
                raise
            sys.exit(1)
        return

    # batch (precisa Playwright)
    if args.mode == "batch":
        if not args.config:
            print("[Erro] --config obrigatório em mode=batch"); sys.exit(1)
        from playwright.sync_api import sync_playwright
        with sync_playwright() as p:
            flow_batch(p, args)
        return

    # Demais modos usam Playwright (list/run/plan/plan-live)
    from playwright.sync_api import sync_playwright
    with sync_playwright() as p:
        # plan-live gera o plano dinâmico e opcionalmente já executa
        if args.mode == "plan-live":
            try:
                cfg = plan_live(p, args)
                Path(args.plan_out).parent.mkdir(parents=True, exist_ok=True)
                Path(args.plan_out).write_text(json.dumps(cfg, ensure_ascii=False, indent=2), encoding="utf-8")
                print(f"[OK] Plano (live) gerado em: {args.plan_out} (combos={len(cfg.get('combos', []))})")
                if args.execute_plan:
                    print("[Exec] Rodando batch com o plano recém-gerado...")
                    args.config = args.plan_out
                    flow_batch(p, args)
            except Exception as e:
                print(f"[ERRO][plan-live] {e}")
                if getattr(args, "plan_verbose", False):
                    raise
                sys.exit(1)
            return

        # Single run/list
        browser = p.chromium.launch(headless=not args.headful, slow_mo=args.slowmo)
        context = browser.new_context(ignore_https_errors=True, viewport={"width": 1366, "height": 900})
        page = context.new_page()
        page.set_default_timeout(60_000)
        page.set_default_navigation_timeout(60_000)
        try:
            if args.preset:
                preset_path = Path(args.presets_file)
                store = load_presets(preset_path)
                pr = (store.get("presets") or {}).get(args.preset)
                if not pr:
                    print(f"[Erro] Preset '{args.preset}' não encontrado em {preset_path}")
                    sys.exit(1)
                for k, v in pr.items():
                    if getattr(args, k, None) in (None, False):
                        setattr(args, k, v)

            if args.mode == "list":
                data = fmt_date(args.data)
                flow_list(
                    context, data, args.secao, args.level,
                    args.key1, args.key1_type, args.key2, args.key2_type,
                    args.out, args.debug_dump,
                    label1=args.label1, label2=args.label2, label3=args.label3
                )
                return

            if args.mode == "run":
                for d in iter_dates(args.data, args.since, args.range):
                    missing = []
                    if not args.key1 or not args.key1_type: missing.append("key1/key1-type")
                    if not args.key2 or not args.key2_type: missing.append("key2/key2-type")
                    if missing:
                        print("[Erro] Parâmetros faltando:", ", ".join(missing)); sys.exit(1)

                    out_path = args.out.replace("{date}", d.replace("/", "-"))
                    bulletin_out = args.bulletin_out.replace("{date}", d.replace("/", "-")) if args.bulletin_out else None

                    flow_run(
                        context, d, args.secao,
                        args.key1, args.key1_type, args.key2, args.key2_type, args.key3, args.key3_type,
                        args.query, args.max_links, out_path, args.debug_dump,
                        args.scrape_detail, args.detail_timeout, args.fallback_date_if_missing,
                        label1=args.label1, label2=args.label2, label3=args.label3,
                        max_scrolls=args.max_scrolls, scroll_pause_ms=args.scroll_pause_ms, stable_rounds=args.stable_rounds,
                        state_file=args.state_file, bulletin=args.bulletin, bulletin_out=bulletin_out
                    )
                return

            if args.mode == "plan":
                if not args.map_file:
                    print("[Erro] --map-file obrigatório em mode=plan"); sys.exit(1)
                # Se você já implementou plan_from_map(), descomente:
                # cfg = plan_from_map(args.map_file, args)
                # Path(args.plan_out).parent.mkdir(parents=True, exist_ok=True)
                # Path(args.plan_out).write_text(json.dumps(cfg, ensure_ascii=False, indent=2), encoding="utf-8")
                # print(f"[OK] Plano (map) gerado em: {args.plan_out} (combos={len(cfg.get('combos', []))})")
                # if args.execute_plan:
                #     args.config = args.plan_out
                #     flow_batch(p, args)
                print("[Aviso] plan_from_map ainda não está finalizado neste arquivo.")
                return

            print(f"[Erro] Modo desconhecido: {args.mode}")
            sys.exit(1)

        finally:
            try:
                browser.close()
            except Exception:
                pass

# ----------------------------------------------------------------------
# Entry point
# ----------------------------------------------------------------------
if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n[Abortado pelo usuário]")
        sys.exit(130)
